#!/usr/bin/env python3
"""Generate a trader-facing Word report for CSI500 model bakeoff results."""

from __future__ import annotations

import json
import sys
from pathlib import Path

import lightgbm as lgb
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from sklearn.ensemble import ExtraTreesRegressor
from sklearn.impute import SimpleImputer
from sklearn.linear_model import Ridge
from sklearn.pipeline import make_pipeline
from sklearn.preprocessing import StandardScaler

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from train_lightgbm import build_category_mappings, build_feature_frame, choose_feature_columns, load_frame


BAKEOFF_DIR = ROOT / "quant_data/csi500_2y_run/model_bakeoff_fast_v2"
OLD_TOPK_DIR = ROOT / "quant_data/csi500_2y_run/topk_tests"
TRAIN_PATH = ROOT / "quant_data/csi500_2y_run/ml_features_ready.parquet"
INFERENCE_PATH = ROOT / "quant_data/csi500_2y_run/inference_features_latest.parquet"
REPORT_PATH = ROOT / "quant_data/csi500_2y_run/model_bakeoff_fast_v2/中证500模型比较_交易员视角报告.docx"

MODEL_NAMES = ["lightgbm_regressor", "lightgbm_ranker", "extra_trees", "ridge"]
TOP_KS = [1, 3, 5, 10, 20]


def fmt_pct(value: float | str | None) -> str:
    try:
        return f"{float(value):.2%}"
    except (TypeError, ValueError):
        return ""


def fmt_num(value: float | str | None, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return ""


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def daily_relevance(df: pd.DataFrame) -> np.ndarray:
    ranks = df.groupby("date", sort=False)["future_return"].rank(pct=True, method="first")
    return np.floor(ranks.fillna(0.0).to_numpy(dtype=np.float32) * 5.0).clip(0, 4).astype(np.int32)


def train_model(model_name: str, X_train: pd.DataFrame, y_return: np.ndarray, y_rank: np.ndarray, group: list[int]):
    if model_name == "lightgbm_regressor":
        model = lgb.LGBMRegressor(
            objective="regression",
            learning_rate=0.04,
            n_estimators=220,
            num_leaves=31,
            subsample=0.9,
            colsample_bytree=0.9,
            random_state=42,
            n_jobs=1,
            verbosity=-1,
        )
        model.fit(X_train, y_return)
        return model
    if model_name == "lightgbm_ranker":
        model = lgb.LGBMRanker(
            objective="lambdarank",
            metric="ndcg",
            learning_rate=0.04,
            n_estimators=180,
            num_leaves=31,
            subsample=0.9,
            colsample_bytree=0.9,
            random_state=42,
            n_jobs=1,
            verbosity=-1,
        )
        model.fit(X_train, y_rank, group=group)
        return model
    if model_name == "extra_trees":
        model = make_pipeline(
            SimpleImputer(strategy="median"),
            ExtraTreesRegressor(
                n_estimators=120,
                max_depth=10,
                min_samples_leaf=20,
                random_state=42,
                n_jobs=-1,
            ),
        )
        model.fit(X_train, y_return)
        return model
    if model_name == "ridge":
        model = make_pipeline(SimpleImputer(strategy="median"), StandardScaler(), Ridge(alpha=20.0))
        model.fit(X_train, y_return)
        return model
    raise ValueError(model_name)


def score_latest_inference() -> pd.DataFrame:
    output_path = BAKEOFF_DIR / "latest_inference_top10_by_model.csv"
    if output_path.exists():
        return pd.read_csv(output_path, dtype={"code": str})

    train = load_frame(TRAIN_PATH)
    infer = load_frame(INFERENCE_PATH)
    train["date"] = pd.to_datetime(train["date"])
    infer["date"] = pd.to_datetime(infer["date"])
    train = train.sort_values("date", kind="stable").reset_index(drop=True)
    feature_cols, categorical_cols = choose_feature_columns(train)
    feature_cols = [col for col in feature_cols if col in infer.columns]
    category_mappings = build_category_mappings(train, infer, categorical_cols)
    X_train = build_feature_frame(train, feature_cols, categorical_cols, category_mappings)
    X_infer = build_feature_frame(infer, feature_cols, categorical_cols, category_mappings)
    y_return = pd.to_numeric(train["future_return"], errors="coerce").fillna(0.0).to_numpy(dtype=np.float32)
    y_rank = daily_relevance(train)
    group = train.groupby("date", sort=False).size().astype(int).tolist()

    latest_rows: list[pd.DataFrame] = []
    keep_cols = [col for col in ["date", "code", "name", "close", "pct_chg", "pct_chg_5d", "pct_chg_20d", "turnover"] if col in infer.columns]
    for model_name in MODEL_NAMES:
        print(f"score latest inference: {model_name}", flush=True)
        model = train_model(model_name, X_train, y_return, y_rank, group)
        scored = infer.loc[:, keep_cols].copy()
        scored["score"] = np.asarray(model.predict(X_infer), dtype=np.float32)
        scored = scored.sort_values("score", ascending=False, kind="mergesort").head(10).copy()
        scored.insert(0, "rank", range(1, len(scored) + 1))
        scored.insert(0, "model", model_name)
        latest_rows.append(scored)

    result = pd.concat(latest_rows, ignore_index=True)
    result.to_csv(output_path, index=False, encoding="utf-8-sig")
    return result


def load_old_model_rows() -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for model in ["all_features", "momentum_liquidity", "valuation_momentum"]:
        for top_k in TOP_KS:
            path = OLD_TOPK_DIR / model / f"topk_{top_k}" / "summary.json"
            if not path.exists():
                continue
            data = json.loads(path.read_text(encoding="utf-8"))
            rows.append(
                {
                    "model": f"old_lightgbm_classifier_{model}",
                    "top_k": top_k,
                    "total_return": data["portfolio_total_return"],
                    "cagr": data["portfolio_cagr"],
                    "max_drawdown": data["portfolio_max_drawdown"],
                    "win_rate": data["portfolio_win_rate"],
                }
            )
    return pd.DataFrame(rows)


def load_old_latest_top10() -> pd.DataFrame:
    rows: list[pd.DataFrame] = []
    for model in ["all_features", "momentum_liquidity", "valuation_momentum"]:
        path = OLD_TOPK_DIR / model / "topk_20" / "oos_predictions.parquet"
        if not path.exists():
            continue
        df = pd.read_parquet(path)
        df["date"] = pd.to_datetime(df["date"])
        latest = df["date"].max()
        scored = df[df["date"] == latest].sort_values("score", ascending=False).head(10).copy()
        scored.insert(0, "rank", range(1, len(scored) + 1))
        scored.insert(0, "model", f"old_lightgbm_classifier_{model}")
        rows.append(scored)
    return pd.concat(rows, ignore_index=True)


def load_existing_live_latest_top10() -> pd.DataFrame:
    rows: list[pd.DataFrame] = []
    live_paths = {
        "old_live_momentum_liquidity": ROOT
        / "quant_data/csi500_2y_run/feature_group_tests/momentum_liquidity/models/inference_scores_latest.parquet",
    }
    for model, path in live_paths.items():
        if not path.exists():
            continue
        df = pd.read_parquet(path)
        df["date"] = pd.to_datetime(df["date"])
        scored = df.sort_values("score", ascending=False, kind="mergesort").head(10).copy()
        scored.insert(0, "rank", range(1, len(scored) + 1))
        scored.insert(0, "model", model)
        rows.append(scored)
    if not rows:
        return pd.DataFrame()
    return pd.concat(rows, ignore_index=True)


def model_plain_name(model: str) -> str:
    return {
        "lightgbm_regressor": "LightGBM Regressor",
        "lightgbm_ranker": "LightGBM Ranker",
        "extra_trees": "ExtraTrees Regressor",
        "ridge": "Ridge Regression",
        "old_lightgbm_classifier_all_features": "旧 LightGBM Classifier - all_features",
        "old_lightgbm_classifier_momentum_liquidity": "旧 LightGBM Classifier - momentum_liquidity",
        "old_lightgbm_classifier_valuation_momentum": "旧 LightGBM Classifier - valuation_momentum",
        "old_live_momentum_liquidity": "原流程最新候选 - momentum_liquidity",
    }.get(model, model)


def main() -> int:
    comparison = pd.read_csv(BAKEOFF_DIR / "model_topk_comparison.csv")
    old_comparison = load_old_model_rows()
    all_comparison = pd.concat(
        [
            comparison.loc[:, ["model", "top_k", "total_return", "cagr", "max_drawdown", "win_rate"]],
            old_comparison,
        ],
        ignore_index=True,
    )
    winners = all_comparison.sort_values(["top_k", "total_return"], ascending=[True, False]).groupby("top_k", as_index=False).first()

    backtest_latest = pd.read_csv(BAKEOFF_DIR / "latest_rebalance_top20_by_model.csv", dtype={"code": str})
    old_latest = load_old_latest_top10()
    latest_inference = score_latest_inference()
    existing_live_latest = load_existing_live_latest_top10()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10)

    doc.add_heading("中证500模型比较报告：交易员视角", level=0)
    doc.add_paragraph("数据范围：500只中证500成分股，228,919行样本，460个交易日。")
    doc.add_paragraph("回测窗口：2025-06-17 至 2026-04-21；调仓频率：每5个交易日；调仓次数：42次。")
    doc.add_paragraph("重要说明：本报告是模型研究结果，不是投资建议；回测未单独加入交易成本、滑点、停牌/涨跌停成交约束。")

    doc.add_heading("Executive Summary", level=1)
    add_bullets(
        doc,
        [
            "如果只选少数强票，LightGBM Regressor 是本轮最强：top1、top3、top5 都胜出，说明直接预测未来收益率比简单判断涨跌更贴近交易目标。",
            "如果想做更分散的组合，ExtraTrees 更适合：top10 和 top20 胜出，胜率更高、回撤也相对温和。",
            "LightGBM Ranker 理论上适合排序，但这次快速版本没有跑赢。原因大概率不是模型类型错，而是排序标签和参数还没精调。",
            "Ridge 是很好的基线模型：速度极快、解释性强，但收益不够强，不建议作为主模型。",
            "你说的 XSBoost 应该是 XGBoost。它和 LightGBM 同属梯度提升树家族，但当前本机环境没有安装 XGBoost，本次没有参与比较。",
            "交易建议：真正当前可看的候选日是 2026-04-30。本报告把 2026-04-30 候选放在前面；2026-04-21 只是回测验证日。",
            "实盘观察优先采用 LightGBM Regressor 的 top5 与 ExtraTrees 的 top10 交集，再用旧 momentum_liquidity 模型做二次确认。",
        ],
    )

    doc.add_heading("跨新旧模型的 top-k 冠军", level=2)
    add_table(
        doc,
        ["top-k", "冠军模型", "总收益", "年化", "最大回撤", "胜率"],
        [
            [
                int(row.top_k),
                model_plain_name(row.model),
                fmt_pct(row.total_return),
                fmt_pct(row.cagr),
                fmt_pct(row.max_drawdown),
                fmt_pct(row.win_rate),
            ]
            for row in winners.itertuples(index=False)
        ],
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("1. 最新可交易候选 top10：2026-04-30", level=1)
    doc.add_paragraph("这一节才是当前最后可用数据日的选股建议。这里没有 future_return，因为未来还没发生；它应该作为观察名单，而不是回测答案。")
    for model in MODEL_NAMES:
        subset = latest_inference[latest_inference["model"] == model].head(10)
        doc.add_heading(model_plain_name(model), level=2)
        add_table(
            doc,
            ["排名", "代码", "名称", "收盘价", "当日涨跌", "5日涨跌", "20日涨跌", "模型分"],
            [
                [
                    int(row.rank),
                    row.code,
                    row.name,
                    fmt_num(row.close, 2),
                    fmt_pct(row.pct_chg / 100.0 if abs(float(row.pct_chg)) > 1 else row.pct_chg),
                    fmt_pct(row.pct_chg_5d),
                    fmt_pct(row.pct_chg_20d),
                    fmt_num(row.score),
                ]
                for row in subset.itertuples(index=False)
            ],
        )

    if not existing_live_latest.empty:
        doc.add_heading("原流程最新候选 top10：2026-04-30", level=2)
        for model in existing_live_latest["model"].unique():
            subset = existing_live_latest[existing_live_latest["model"] == model].head(10)
            doc.add_heading(model_plain_name(model), level=3)
            add_table(
                doc,
                ["排名", "代码", "名称", "收盘价", "当日涨跌", "5日涨跌", "20日涨跌", "模型分"],
                [
                    [
                        int(row.rank),
                        row.code,
                        row.name,
                        fmt_num(row.close, 2),
                        fmt_pct(row.pct_chg / 100.0 if abs(float(row.pct_chg)) > 1 else row.pct_chg),
                        fmt_pct(row.pct_chg_5d),
                        fmt_pct(row.pct_chg_20d),
                        fmt_num(row.score),
                    ]
                    for row in subset.itertuples(index=False)
                ],
            )

    doc.add_heading("2. 从交易员角度理解这些模型", level=1)
    add_table(
        doc,
        ["模型", "交易员视角", "适合用途", "主要风险"],
        [
            ["旧 LightGBM Classifier", "先判断未来是否上涨，再按上涨概率排序。", "稳健候选池、信号过滤。", "概率高不等于收益高，可能错过弹性票。"],
            ["LightGBM Regressor", "直接预测未来收益率，谁预计涨得多排前面。", "top1/top3/top5 强票。", "更尖锐，也更容易集中到高波动票。"],
            ["LightGBM Ranker", "只关心每日股票之间谁排前面。", "理论上最贴近选股排序。", "需要认真设计排序标签；快速版表现一般。"],
            ["ExtraTrees", "很多随机树投票，结果通常更平滑。", "top10/top20 分散组合。", "单票爆发力不如 Regressor。"],
            ["Ridge", "线性打分，像把因子按固定权重加总。", "基线、解释、检验特征有效性。", "难捕捉非线性关系。"],
            ["XGBoost", "和 LightGBM 类似的梯度提升树。", "可作为下一个对照模型。", "本机未安装，速度通常略慢于 LightGBM。"],
        ],
    )

    doc.add_heading("3. 新四模型完整表现", level=1)
    for top_k in TOP_KS:
        subset = comparison[comparison["top_k"] == top_k].sort_values("total_return", ascending=False)
        doc.add_heading(f"top{top_k} 排名", level=2)
        add_table(
            doc,
            ["模型", "总收益", "年化", "最大回撤", "胜率"],
            [
                [model_plain_name(row.model), fmt_pct(row.total_return), fmt_pct(row.cagr), fmt_pct(row.max_drawdown), fmt_pct(row.win_rate)]
                for row in subset.itertuples(index=False)
            ],
        )

    doc.add_heading("4. 交易解读与组合建议", level=1)
    add_bullets(
        doc,
        [
            "进攻型：使用 LightGBM Regressor 的 top3 或 top5；适合你想抓强票，但仓位要小，且必须人工复核事件和流动性。",
            "均衡型：使用 ExtraTrees 的 top10；它在本次回测里胜率最高，适合作为候选池的主列表。",
            "确认型：用旧 momentum_liquidity 与新模型交叉。两个模型都靠前的票，比单一模型靠前更值得看。",
            "不建议：直接单押 top1。top1 回测很漂亮，但42次调仓样本太少，容易被少数大行情放大。",
            "建议执行流程：先取 LightGBM Regressor top10，再取 ExtraTrees top20，交集进入重点池；旧 momentum_liquidity 排名前30的留下，其他降权。",
        ],
    )

    doc.add_heading("5. 回测最后一次调仓 top10：2026-04-21", level=1)
    doc.add_paragraph("这一节用于评价模型，因为可以看到后续 future_return。不要把它当成今天的实时候选。")
    for model in MODEL_NAMES:
        subset = backtest_latest[backtest_latest["model"] == model].head(10)
        doc.add_heading(model_plain_name(model), level=2)
        add_table(
            doc,
            ["排名", "代码", "名称", "未来收益", "模型分"],
            [
                [i + 1, row.code, row.name, fmt_pct(row.future_return), fmt_num(row.score)]
                for i, row in enumerate(subset.itertuples(index=False))
            ],
        )

    doc.add_heading("6. 旧 LightGBM 模型最后一次调仓 top10：2026-04-21", level=1)
    for model in old_latest["model"].unique():
        subset = old_latest[old_latest["model"] == model].head(10)
        doc.add_heading(model_plain_name(model), level=2)
        add_table(
            doc,
            ["排名", "代码", "名称", "未来收益", "模型分"],
            [
                [int(row.rank), row.code, row.name, fmt_pct(row.future_return), fmt_num(row.score)]
                for row in subset.itertuples(index=False)
            ],
        )

    doc.add_heading("7. 风险与下一步", level=1)
    add_bullets(
        doc,
        [
            "样本量风险：本次 OOS 只有42次调仓，适合做方向判断，还不能单独证明可实盘。",
            "成本风险：当前回测没有显式计入佣金、滑点、冲击成本；top1/top3 对这些尤其敏感。",
            "成交风险：A股涨跌停、停牌、开盘跳空会影响实际成交价。",
            "下一步建议：做融合模型，把 LightGBM Regressor 排名、ExtraTrees 排名、旧 momentum_liquidity 排名加权，观察是否能同时保留收益和胜率。",
        ],
    )

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
