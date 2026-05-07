from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "quant_data" / "comparison_reports"
OUT_PATH = OUT_DIR / "trader_view_sse50_vs_csi500_2y_report.docx"


def read_json(rel_path: str) -> dict:
    return json.loads((ROOT / rel_path).read_text(encoding="utf-8"))


def pct(value: float) -> str:
    return f"{float(value) * 100:.1f}%"


def f2(value: float) -> str:
    return f"{float(value):.2f}"


def f3(value: float) -> str:
    return f"{float(value):.3f}"


def diagnostics(run_name: str, group: str) -> dict:
    return read_json(f"quant_data/{run_name}/diagnostics/{group}/diagnostics_summary.json")


def feature_group_table(run_name: str) -> pd.DataFrame:
    payload = read_json(f"quant_data/{run_name}/feature_group_tests/group_run_results.json")
    rows = []
    for name, item in payload.items():
        if not item.get("ok"):
            continue
        summary = item["summary"]
        rows.append(
            {
                "信号组合": name,
                "总收益": pct(summary["portfolio_total_return"]),
                "年化": pct(summary["portfolio_cagr"]),
                "最大回撤": pct(summary["portfolio_max_drawdown"]),
                "胜率": pct(summary["portfolio_win_rate"]),
                "每期平均收益": pct(summary["portfolio_avg_return"]),
                "_sort": summary["portfolio_total_return"],
            }
        )
    return pd.DataFrame(rows).sort_values("_sort", ascending=False).drop(columns=["_sort"])


def topk_table(run_name: str, n: int = 12) -> pd.DataFrame:
    df = pd.read_csv(ROOT / f"quant_data/{run_name}/topk_tests/topk_results.csv")
    df = df[df["ok"] == True].copy()
    df["收益回撤比"] = df["total_return"] / df["max_drawdown"].abs()
    df = df.sort_values("total_return", ascending=False).head(n)
    return pd.DataFrame(
        {
            "信号组合": df["group"],
            "持仓数": df["top_k"].astype(int),
            "总收益": df["total_return"].map(pct),
            "年化": df["cagr"].map(pct),
            "最大回撤": df["max_drawdown"].map(pct),
            "胜率": df["win_rate"].map(pct),
            "每期波动": df["std_return"].map(pct),
            "收益回撤比": df["收益回撤比"].map(lambda x: f"{x:.2f}"),
        }
    )


def latest_top10(path: str) -> pd.DataFrame:
    df = pd.read_parquet(ROOT / path).sort_values("score", ascending=False).head(10).copy()
    cols = ["code", "name", "score"]
    return pd.DataFrame(
        {
            "排名": range(1, len(df) + 1),
            "代码": df["code"].astype(str).str.zfill(6),
            "名称": df["name"].astype(str),
            "分数": df["score"].map(lambda x: f"{x:.4f}"),
        }
    )


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(20)


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Microsoft YaHei"
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
    doc.add_paragraph("")


def add_df(doc: Document, df: pd.DataFrame) -> None:
    add_table(doc, list(df.columns), df.astype(str).values.tolist())


def bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_paragraphs(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    sse_run = read_json("quant_data/sse50_2y_run/run_meta.json")
    csi_run = read_json("quant_data/csi500_2y_run/run_meta.json")
    sse_feat = read_json("quant_data/sse50_2y_run/ml_features_ready.parquet.meta.json")
    csi_feat = read_json("quant_data/csi500_2y_run/ml_features_ready.parquet.meta.json")
    sse_base = read_json("quant_data/sse50_2y_run/backtest/summary.json")
    csi_base = read_json("quant_data/csi500_2y_run/backtest/summary.json")
    sse_diag = diagnostics("sse50_2y_run", "valuation_momentum")
    csi_diag_ml = diagnostics("csi500_2y_run", "momentum_liquidity")
    csi_diag_vm = diagnostics("csi500_2y_run", "valuation_momentum")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        doc.styles[style_name].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "交易员视角：上证50 vs 中证500 两年选股实验报告")
    doc.add_paragraph("生成日期：2026-04-30")
    doc.add_paragraph(
        "这版报告弱化模型术语，重点讨论交易含义：哪个股票池更值得做、收益来自哪里、持仓怎么选、"
        "有没有明显 beta、流动性信号是否可用、回撤和成本风险在哪里。"
    )

    doc.add_heading("第 1 页：给交易员的结论", level=1)
    add_paragraphs(
        doc,
        [
            "一句话：如果把这套信号当作一个交易想法，中证500比上证50更值得继续推进。",
            "上证50也能做出收益，但有效性主要集中在最靠前的少数股票，持仓越集中越漂亮，也越危险。中证500的股票池更宽，信号分层更顺，Top5/Top10/Top15 都能跑出较好的收益，交易上更容易做成组合，而不是押单票。",
            "最重要的发现是：上证50里成交和流动性类信号没有太大帮助；中证500里成交和流动性类信号明显有价值。这符合交易直觉：上证50都是大票，流动性差异本来就小；中证500里资金关注度和量能变化更能区分强弱。",
        ],
    )
    bullet(
        doc,
        [
            f"上证50最佳交易候选：valuation_momentum Top5，组合总收益 {pct(sse_diag['benchmark']['strategy_top5_total_return'])}，等权基准 {pct(sse_diag['benchmark']['equal_weight_total_return'])}。",
            f"中证500最佳交易候选：momentum_liquidity Top5，组合总收益 {pct(csi_diag_ml['benchmark']['strategy_top5_total_return'])}，等权基准 {pct(csi_diag_ml['benchmark']['equal_weight_total_return'])}。",
            "中证500 Top5 的相对等权权益超额约 77.7%，这是本次最像 alpha 的地方。",
            "但这还不是实盘结论：没有交易成本，没有历史动态成分股，样本只有42次调仓，仍需严谨复核。",
        ],
    )

    doc.add_heading("第 2 页：这次实验到底模拟了什么交易", level=1)
    bullet(
        doc,
        [
            "每5个交易日调仓一次。",
            "每次在股票池内按信号分数排序。",
            "买入排名靠前的 TopK 股票，组合内等权。",
            "收益用未来5个交易日的股票收益衡量。",
            "上证50测试了50只股票，中证500测试了500只股票。",
            "回测期真正发生交易的区间为 2025-06-17 到 2026-04-21，共42次调仓。",
        ],
    )
    doc.add_paragraph("从交易角度看，这不是预测指数涨跌，而是在同一个股票池里做相对强弱选择。")

    doc.add_heading("第 3 页：两个股票池的交易性格不同", level=1)
    add_table(
        doc,
        ["维度", "上证50", "中证500", "交易含义"],
        [
            ["股票数量", sse_run["stock_count"], csi_run["stock_count"], "中证500更有横截面选择空间"],
            ["训练样本", f"{sse_feat['train_rows']:,}", f"{csi_feat['train_rows']:,}", "样本越多，信号稳定性越容易评估"],
            ["股票类型", "大盘核心资产", "中盘成长/周期/制造更多", "中证500风格更分散"],
            ["流动性差异", "较小", "更大", "成交/换手信号在中证500更有用"],
            ["组合风险", "单票影响更大", "可分散", "中证500更适合做组合"],
        ],
    )

    doc.add_heading("第 4 页：核心收益对比", level=1)
    add_table(
        doc,
        ["策略", "总收益", "年化", "最大回撤", "胜率", "每期平均收益"],
        [
            ["上证50 全特征 Top5", pct(sse_base["portfolio_total_return"]), pct(sse_base["portfolio_cagr"]), pct(sse_base["portfolio_max_drawdown"]), pct(sse_base["portfolio_win_rate"]), pct(sse_base["portfolio_avg_return"])],
            ["上证50 valuation_momentum Top5", pct(sse_diag["benchmark"]["strategy_top5_total_return"]), "见回测表", "-8.0%左右", "约59.5%", "约1.0%"],
            ["中证500 全特征 Top5", pct(csi_base["portfolio_total_return"]), pct(csi_base["portfolio_cagr"]), pct(csi_base["portfolio_max_drawdown"]), pct(csi_base["portfolio_win_rate"]), pct(csi_base["portfolio_avg_return"])],
            ["中证500 momentum_liquidity Top5", pct(csi_diag_ml["benchmark"]["strategy_top5_total_return"]), "约195.7%", "-11.0%左右", "61.9%", "约2.37%"],
        ],
    )
    doc.add_paragraph("中证500不仅收益更高，回撤并没有明显恶化，这是最值得继续研究的地方。")

    doc.add_heading("第 5 页：跑赢的是 beta 还是 alpha", level=1)
    add_table(
        doc,
        ["组合", "策略收益", "等权基准", "相对等权权益超额", "超额为正比例"],
        [
            ["上证50 valuation_momentum Top5", pct(sse_diag["benchmark"]["strategy_top5_total_return"]), pct(sse_diag["benchmark"]["equal_weight_total_return"]), pct(sse_diag["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"]), pct(sse_diag["benchmark"]["excess_win_rate"])],
            ["中证500 valuation_momentum Top5", pct(csi_diag_vm["benchmark"]["strategy_top5_total_return"]), pct(csi_diag_vm["benchmark"]["equal_weight_total_return"]), pct(csi_diag_vm["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"]), pct(csi_diag_vm["benchmark"]["excess_win_rate"])],
            ["中证500 momentum_liquidity Top5", pct(csi_diag_ml["benchmark"]["strategy_top5_total_return"]), pct(csi_diag_ml["benchmark"]["equal_weight_total_return"]), pct(csi_diag_ml["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"]), pct(csi_diag_ml["benchmark"]["excess_win_rate"])],
        ],
    )
    doc.add_paragraph(
        "等权基准本身也涨了，说明市场环境有 beta 顺风；但策略显著跑赢等权，尤其是中证500 momentum_liquidity，"
        "说明不只是吃市场上涨，也存在股票选择贡献。是否是真 alpha，还需要扣成本、做动态成分股、做行业市值暴露拆解。"
    )

    doc.add_heading("第 6 页：什么信号在赚钱", level=1)
    doc.add_heading("上证50信号组合排名", level=2)
    add_df(doc, feature_group_table("sse50_2y_run"))
    doc.add_heading("中证500信号组合排名", level=2)
    add_df(doc, feature_group_table("csi500_2y_run"))
    add_paragraphs(
        doc,
        [
            "上证50最好的组合是估值加动量，说明在大票里，价格相对位置和估值状态比较关键。",
            "中证500最好的组合是动量加流动性，说明在中盘股里，量能和换手变化更像交易资金的脚印。",
            "这对交易有启发：不要把同一套因子机械套到所有股票池。大票和中盘票的有效信号可能不同。",
        ],
    )

    doc.add_heading("第 7 页：持仓数 TopK 怎么选", level=1)
    doc.add_heading("上证50 TopK 排名", level=2)
    add_df(doc, topk_table("sse50_2y_run"))
    doc.add_heading("中证500 TopK 排名", level=2)
    add_df(doc, topk_table("csi500_2y_run"))
    add_paragraphs(
        doc,
        [
            "上证50里 Top1/Top3 很漂亮，但交易上不舒服，因为单票路径依赖太强。上证50如果做，Top5比Top1更像可交易组合。",
            "中证500里 Top1 并不稳定，Top5 是本次最优，Top10/Top15 也不差。交易上我更愿意从 Top5 到 Top15 做容量和成本测试。",
            "如果实盘资金较小，可以研究 Top5；如果资金更大，Top10/Top15 可能更稳。"
        ],
    )

    doc.add_heading("第 8 页：分层是否像一个可交易信号", level=1)
    doc.add_heading("上证50 valuation_momentum 五分组", level=2)
    add_df(doc, pd.DataFrame(sse_diag["q5_stats"]))
    doc.add_heading("中证500 momentum_liquidity 五分组", level=2)
    add_df(doc, pd.DataFrame(csi_diag_ml["q5_stats"]))
    doc.add_paragraph(
        "交易员看分层，重点不是模型术语，而是最高分组是否明显强于低分组。中证500这点更好：Q1到Q5大体递增，"
        "说明分数更像可交易排序。上证50是Q5最好，但中间层比较乱，更像顶部有效、整体排序一般。"
    )

    doc.add_heading("第 9 页：回撤和路径风险", level=1)
    bullet(
        doc,
        [
            "上证50 Top1 虽然收益高，但单票集中，路径风险最大。",
            "中证500 momentum_liquidity Top5 收益最高且最大回撤约-11%，在本次样本里收益回撤比很突出。",
            "中证500 Top10/Top15 收益略低但持仓更分散，更可能承载资金。",
            "当前回测未加交易成本，真实回撤可能更大，收益可能被压低。",
        ],
    )

    doc.add_heading("第 10 页：交易成本和换手风险", level=1)
    add_paragraphs(
        doc,
        [
            "这类策略每5个交易日调仓一次，成本一定要严肃处理。即使是中证500，Top5组合如果每期换手高，手续费、印花税、滑点都会明显吃收益。",
            "成本压力的排序大致是：Top1最高、Top5较高、Top10/Top15更可控。因为持仓越少，单票进出权重越大，价格冲击也越明显。",
            "下一版回测应该输出每期换手率、单票成交金额、成本后收益、成本后最大回撤。没有成本前，不建议把收益数字当作实盘预期。"
        ],
    )

    doc.add_heading("第 11 页：最新中证500交易观察名单", level=1)
    doc.add_paragraph("基于中证500 momentum_liquidity 模型，对2026-04-30最新截面打分，Top10如下：")
    add_df(doc, latest_top10("quant_data/csi500_2y_run/feature_group_tests/momentum_liquidity/models/inference_scores_latest.parquet"))
    doc.add_paragraph("这不是买入建议，只是当前模型输出的观察名单。实盘前至少要叠加交易成本、停牌/涨跌停、成交额、风险敞口和人工复核。")

    doc.add_heading("第 12 页：适合怎样的交易使用方式", level=1)
    bullet(
        doc,
        [
            "更适合作为周频/五日频的股票池打分器，而不是日内交易工具。",
            "更适合中证500这类较宽股票池，而不是只在50只大票里硬挑。",
            "适合和人工判断结合：模型给候选池，交易员再看事件、行业、成交、风险。",
            "不适合直接单票满仓，因为Top1路径风险过大。",
        ],
    )

    doc.add_heading("第 13 页：最需要补的风控", level=1)
    bullet(
        doc,
        [
            "单票最大权重限制，例如不超过20%。",
            "行业最大权重限制，避免集中在同一个行业。",
            "涨跌停过滤，避免无法成交。",
            "成交额过滤，避免模型选到容量不足的票。",
            "止损或降仓规则，尤其是连续两期超额为负时。",
            "市场状态过滤，例如指数大跌或流动性收缩时降低仓位。",
        ],
    )

    doc.add_heading("第 14 页：我会怎么推进到下一轮", level=1)
    bullet(
        doc,
        [
            "第一步：给中证500 momentum_liquidity Top5/10/15 加交易成本。",
            "第二步：计算每期换手率和实际成交容量。",
            "第三步：做行业、市值、波动率暴露拆解，判断是不是某个风格在贡献。",
            "第四步：使用历史动态成分股，减少幸存者偏差。",
            "第五步：保留一个完全不碰的测试期，避免过度挑参数。",
        ],
    )

    doc.add_heading("第 15 页：最终交易员版判断", level=1)
    add_paragraphs(
        doc,
        [
            "这不是一个已经可以直接实盘的策略，但它已经从“模型实验”进入“值得交易化验证”的阶段。",
            "上证50方向更像小样本信号探索，中证500方向更像可组合化交易研究。中证500里成交/换手和动量结合的信号表现最好，且分层较健康，这是本次最重要的交易发现。",
            "下一步的胜负手不是继续调模型，而是把交易现实补上：成本、换手、容量、风格暴露、动态成分股、独立测试期。只有这些做完，才能判断它到底是可交易 alpha，还是一段历史里的漂亮回测。"
        ],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
