#!/usr/bin/env python3
"""Run an index-universe research pipeline and generate a trader-oriented report."""

from __future__ import annotations

import argparse
import json
import math
import shutil
import subprocess
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import download_data as dd

BASE_FEATURES = [
    "open",
    "high",
    "low",
    "close",
    "volume",
    "amount",
    "turnover",
    "amplitude",
    "pct_chg",
    "change",
    "close_val",
    "pct_chg_val",
    "pe_ttm",
    "pb",
    "ps",
    "pcf",
    "ma5",
    "ma20",
    "bias_20",
    "pct_chg_5d",
    "pct_chg_20d",
    "volatility_20d",
    "turnover_ma5",
    "volume_ma5",
    "close_to_high_20d",
    "close_to_low_20d",
]

FEATURE_GROUPS: dict[str, list[str]] = {
    "all_features": BASE_FEATURES,
    "momentum_trend": [
        "ma5",
        "ma20",
        "bias_20",
        "pct_chg_5d",
        "pct_chg_20d",
        "close_to_high_20d",
        "close_to_low_20d",
        "pct_chg",
        "change",
    ],
    "valuation": ["pe_ttm", "pb", "ps", "pcf"],
    "liquidity_volume": ["volume", "amount", "turnover", "turnover_ma5", "volume_ma5"],
    "volatility_position": ["volatility_20d", "amplitude", "close_to_high_20d", "close_to_low_20d"],
    "price_ohlc": ["open", "high", "low", "close", "close_val"],
    "valuation_momentum": [
        "pe_ttm",
        "pb",
        "ps",
        "pcf",
        "ma5",
        "ma20",
        "bias_20",
        "pct_chg_5d",
        "pct_chg_20d",
        "close_to_high_20d",
        "close_to_low_20d",
        "pct_chg",
        "change",
    ],
    "momentum_liquidity": [
        "ma5",
        "ma20",
        "bias_20",
        "pct_chg_5d",
        "pct_chg_20d",
        "pct_chg",
        "change",
        "volume",
        "amount",
        "turnover",
        "turnover_ma5",
        "volume_ma5",
    ],
}

TOPK_GROUPS = ["all_features", "valuation_momentum", "momentum_liquidity"]
TOPK_VALUES = [1, 3, 5, 10, 15, 20, 30]
DIAGNOSTIC_GROUPS = ["valuation_momentum", "momentum_liquidity"]


def pct(value: Any) -> str:
    try:
        if value is None or (isinstance(value, float) and math.isnan(value)):
            return "N/A"
        return f"{float(value) * 100:.1f}%"
    except Exception:
        return "N/A"


def f3(value: Any) -> str:
    try:
        if value is None or (isinstance(value, float) and math.isnan(value)):
            return "N/A"
        return f"{float(value):.3f}"
    except Exception:
        return "N/A"


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def normalize_csindex_constituents(df: pd.DataFrame, universe: str) -> tuple[pd.DataFrame, str | None]:
    exchange_map = {
        "深圳证券交易所": "sz",
        "Shenzhen Stock Exchange": "sz",
        "上海证券交易所": "sh",
        "Shanghai Stock Exchange": "sh",
    }
    out = pd.DataFrame(
        {
            "code": df["成分券代码"].astype(str).str.zfill(6),
            "exchange": df["交易所"].map(exchange_map).fillna(df.get("交易所英文名称", "").map(exchange_map)),
            "name": df["成分券名称"].astype(str),
            "industry": "UNKNOWN",
            "industry_classification": "UNKNOWN",
            "update_date": pd.to_datetime(df["日期"], errors="coerce"),
            "universe": universe,
        }
    )
    out["exchange"] = out["exchange"].fillna(out["code"].map(lambda x: "sh" if x.startswith(("5", "6", "9")) else "sz"))
    out = out.drop_duplicates(subset=["code"]).sort_values("code").reset_index(drop=True)
    date_max = None
    if out["update_date"].notna().any():
        date_max = str(out["update_date"].max().date())
    return out, date_max


def build_basic_valuation_df(bundle_df: pd.DataFrame, code: str) -> pd.DataFrame:
    df = bundle_df.copy()
    df["code"] = code
    df["pct_chg"] = df["pctChg"]
    df = df.rename(columns={"peTTM": "pe_ttm", "pbMRQ": "pb", "psTTM": "ps", "pcfNcfTTM": "pcf"})
    cols = ["date", "code", "exchange", "close", "pct_chg", "pe_ttm", "pb", "ps", "pcf"]
    return df[[col for col in cols if col in df.columns]].copy()


def download_index_data(
    *,
    run_dir: Path,
    universe: str,
    index_code: str,
    start_date: str,
    end_date: str,
    sleep: float,
    overwrite: bool,
) -> None:
    dd.load_dependencies()
    import akshare as ak

    run_dir.mkdir(parents=True, exist_ok=True)
    kline_dir, valuation_dir = dd.ensure_dirs(run_dir)

    print(f"[step1] fetching constituents: {index_code}", flush=True)
    cons = ak.index_stock_cons_csindex(symbol=index_code)
    stock_df, date_max = normalize_csindex_constituents(cons, universe)
    stock_df.to_parquet(run_dir / "stock_list.parquet", index=False)
    stock_df.to_csv(run_dir / "stock_list.csv", index=False, encoding="utf-8-sig")

    failures: list[dict[str, str]] = []
    success = 0
    dd.baostock_login()
    started = time.time()
    try:
        for idx, row in enumerate(stock_df.itertuples(index=False), start=1):
            code = str(row.code).zfill(6)
            exchange = str(row.exchange).lower()
            kline_path = kline_dir / f"{code}.parquet"
            valuation_path = valuation_dir / f"{code}.parquet"
            if not overwrite and kline_path.exists() and valuation_path.exists():
                success += 1
                continue

            bundle_df, reason = dd.download_baostock_daily_bundle(
                code,
                exchange=exchange,
                start_date=start_date,
                end_date=end_date,
            )
            if bundle_df is None:
                failures.append({"code": code, "exchange": exchange, "reason": reason or "unknown"})
                print(f"[step1] {idx}/{len(stock_df)} {code} failed: {reason}", flush=True)
                time.sleep(sleep)
                continue

            dd.build_kline_df(bundle_df, code).to_parquet(kline_path, index=False)
            build_basic_valuation_df(bundle_df, code).to_parquet(valuation_path, index=False)
            success += 1
            if idx % 100 == 0 or idx == len(stock_df):
                print(f"[step1] downloaded {idx}/{len(stock_df)} success={success}", flush=True)
            time.sleep(sleep)
    finally:
        dd.baostock_logout()

    if failures:
        pd.DataFrame(failures).to_csv(run_dir / "download_failures.csv", index=False, encoding="utf-8-sig")

    summary = pd.DataFrame(
        [
            {
                "universe": universe,
                "index_code": index_code,
                "stock_count": len(stock_df),
                "download_success_count": success,
                "download_failure_count": len(failures),
                "start_date": start_date,
                "end_date": end_date,
                "elapsed_seconds": round(time.time() - started, 2),
            }
        ]
    )
    summary.to_csv(run_dir / "download_summary.csv", index=False, encoding="utf-8-sig")
    write_json(
        run_dir / "run_meta.json",
        {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "universe": universe,
            "index_code": index_code,
            "constituent_source": f"akshare.index_stock_cons_csindex({index_code})",
            "constituent_date_max": date_max,
            "start_date": start_date,
            "end_date": end_date,
            "stock_count": int(len(stock_df)),
            "download_success_count": int(success),
            "download_failure_count": int(len(failures)),
            "elapsed_seconds": round(time.time() - started, 2),
        },
    )


def run_py(args: list[str]) -> None:
    print("+ " + " ".join(args), flush=True)
    subprocess.run([sys.executable, *args], cwd=ROOT, check=True)


def make_feature_group_files(run_dir: Path) -> None:
    src = pd.read_parquet(run_dir / "ml_features_ready.parquet")
    out_root = run_dir / "feature_group_tests"
    out_root.mkdir(parents=True, exist_ok=True)
    manifest: dict[str, Any] = {}
    meta_cols = ["date", "code", "exchange", "name", "industry", "future_return", "label"]
    available = set(src.columns)
    for group, features in FEATURE_GROUPS.items():
        selected_features = [col for col in features if col in available]
        cols = [col for col in meta_cols if col in available] + selected_features
        out_dir = out_root / group
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / "features.parquet"
        src[cols].to_parquet(out_path, index=False)
        manifest[group] = {
            "path": str(out_path),
            "rows": int(len(src)),
            "cols": int(len(cols)),
            "features": selected_features,
            "feature_count": int(len(selected_features)),
        }
    write_json(out_root / "manifest.json", manifest)


def run_feature_group_backtests(run_dir: Path, *, label_horizon: int, label_threshold: float) -> None:
    results: dict[str, Any] = {}
    for group in FEATURE_GROUPS:
        start = time.time()
        out_dir = run_dir / "feature_group_tests" / group / "backtest"
        try:
            run_py(
                [
                    "backtest_walk_forward.py",
                    "--train-path",
                    str(run_dir / "feature_group_tests" / group / "features.parquet"),
                    "--output-dir",
                    str(out_dir),
                    "--min-train-days",
                    "252",
                    "--retrain-every",
                    "20",
                    "--rebalance-every",
                    "5",
                    "--top-k",
                    "5",
                    "--threshold",
                    "0.5",
                    "--profile-name",
                    group,
                    "--profile-label",
                    group,
                    "--label-horizon",
                    str(label_horizon),
                    "--label-threshold",
                    str(label_threshold),
                ]
            )
            results[group] = {
                "ok": True,
                "elapsed_seconds": round(time.time() - start, 2),
                "summary": read_json(out_dir / "summary.json"),
            }
        except Exception as exc:
            results[group] = {"ok": False, "elapsed_seconds": round(time.time() - start, 2), "error": str(exc)}
    write_json(run_dir / "feature_group_tests" / "group_run_results.json", results)


def run_topk_tests(run_dir: Path, *, label_horizon: int, label_threshold: float) -> None:
    rows: list[dict[str, Any]] = []
    out_root = run_dir / "topk_tests"
    for group in TOPK_GROUPS:
        for top_k in TOPK_VALUES:
            start = time.time()
            out_dir = out_root / group / f"topk_{top_k}"
            row: dict[str, Any] = {"group": group, "top_k": top_k}
            try:
                run_py(
                    [
                        "backtest_walk_forward.py",
                        "--train-path",
                        str(run_dir / "feature_group_tests" / group / "features.parquet"),
                        "--output-dir",
                        str(out_dir),
                        "--min-train-days",
                        "252",
                        "--retrain-every",
                        "20",
                        "--rebalance-every",
                        "5",
                        "--top-k",
                        str(top_k),
                        "--threshold",
                        "0.5",
                        "--profile-name",
                        group,
                        "--profile-label",
                        group,
                        "--label-horizon",
                        str(label_horizon),
                        "--label-threshold",
                        str(label_threshold),
                    ]
                )
                summary = read_json(out_dir / "summary.json")
                row.update(
                    {
                        "elapsed_seconds": round(time.time() - start, 2),
                        "ok": True,
                        "total_return": summary["portfolio_total_return"],
                        "cagr": summary["portfolio_cagr"],
                        "max_drawdown": summary["portfolio_max_drawdown"],
                        "win_rate": summary["portfolio_win_rate"],
                        "avg_return": summary["portfolio_avg_return"],
                        "std_return": summary["portfolio_std_return"],
                        "auc": summary["oos_metrics"]["auc"],
                        "precision": summary["oos_metrics"]["precision"],
                        "recall": summary["oos_metrics"]["recall"],
                        "num_rebalances": summary["num_rebalances"],
                    }
                )
            except Exception as exc:
                row.update({"elapsed_seconds": round(time.time() - start, 2), "ok": False, "error": str(exc)})
            rows.append(row)
    out_root.mkdir(parents=True, exist_ok=True)
    df = pd.DataFrame(rows)
    df.to_csv(out_root / "topk_results.csv", index=False, encoding="utf-8-sig")
    write_json(out_root / "topk_results.json", {"rows": rows})


def equity_stats(period_returns: pd.Series) -> dict[str, Any]:
    returns = pd.to_numeric(period_returns, errors="coerce").dropna()
    if returns.empty:
        return {"avg_period_return": None, "win_rate": None, "total_return": None, "max_drawdown": None, "periods": 0}
    equity = (1.0 + returns).cumprod()
    drawdown = equity / equity.cummax() - 1.0
    return {
        "avg_period_return": float(returns.mean()),
        "win_rate": float((returns > 0).mean()),
        "total_return": float(equity.iloc[-1] - 1.0),
        "max_drawdown": float(drawdown.min()),
        "periods": int(len(returns)),
    }


def quantile_returns(pred: pd.DataFrame, q: int) -> tuple[pd.DataFrame, pd.DataFrame]:
    rows: list[dict[str, Any]] = []
    for date, group in pred.groupby("date"):
        ranked = group.sort_values("score", ascending=True).copy()
        bucket = pd.qcut(np.arange(len(ranked)), q=q, labels=[f"Q{i}" for i in range(1, q + 1)])
        ranked["bucket"] = bucket.astype(str)
        for b, bdf in ranked.groupby("bucket"):
            rows.append({"rebalance_date": date, "bucket": b, "return": float(bdf["future_return"].mean())})
    curve = pd.DataFrame(rows)
    stats = []
    for bucket, bdf in curve.groupby("bucket"):
        item = {"bucket": bucket, **equity_stats(bdf["return"])}
        stats.append(item)
    high = curve[curve["bucket"] == f"Q{q}"].set_index("rebalance_date")["return"]
    low = curve[curve["bucket"] == "Q1"].set_index("rebalance_date")["return"]
    spread = (high - low).dropna().reset_index(name="return")
    stats.append({"bucket": f"Q{q}-Q1", **equity_stats(spread["return"])})
    return curve, pd.DataFrame(stats)


def run_diagnostics(run_dir: Path) -> None:
    for group in DIAGNOSTIC_GROUPS:
        source = run_dir / "feature_group_tests" / group / "backtest"
        pred = pd.read_parquet(source / "oos_predictions.parquet")
        trade = pd.read_parquet(source / "trade_log.parquet")
        pred["date"] = pd.to_datetime(pred["date"])
        trade["rebalance_date"] = pd.to_datetime(trade["rebalance_date"])

        ic_rows = []
        eq_rows = []
        for date, gdf in pred.groupby("date"):
            ic = gdf["score"].corr(gdf["future_return"], method="spearman")
            ic_rows.append({"rebalance_date": date, "rank_ic": ic, "stocks": len(gdf)})
            top_ret = trade.loc[trade["rebalance_date"] == date, "future_return"].mean()
            ew_ret = gdf["future_return"].mean()
            eq_rows.append(
                {
                    "rebalance_date": date,
                    "strategy_return": top_ret,
                    "equal_weight_return": ew_ret,
                    "excess_return": top_ret - ew_ret,
                }
            )

        ic_df = pd.DataFrame(ic_rows).dropna(subset=["rank_ic"])
        eq_df = pd.DataFrame(eq_rows).dropna()
        eq_df["strategy_equity"] = (1.0 + eq_df["strategy_return"]).cumprod()
        eq_df["equal_weight_equity"] = (1.0 + eq_df["equal_weight_return"]).cumprod()
        eq_df["excess_equity_ratio"] = eq_df["strategy_equity"] / eq_df["equal_weight_equity"]

        q5_curve, q5_stats = quantile_returns(pred, 5)
        q10_curve, q10_stats = quantile_returns(pred, 10)

        out = run_dir / "diagnostics" / group
        out.mkdir(parents=True, exist_ok=True)
        ic_df.to_csv(out / "rank_ic_timeseries.csv", index=False, encoding="utf-8-sig")
        eq_df.to_csv(out / "benchmark_excess.csv", index=False, encoding="utf-8-sig")
        q5_curve.to_csv(out / "quantile_5_curve.csv", index=False, encoding="utf-8-sig")
        q5_stats.to_csv(out / "quantile_5_stats.csv", index=False, encoding="utf-8-sig")
        q10_curve.to_csv(out / "quantile_10_curve.csv", index=False, encoding="utf-8-sig")
        q10_stats.to_csv(out / "quantile_10_stats.csv", index=False, encoding="utf-8-sig")

        ic_std = float(ic_df["rank_ic"].std(ddof=0)) if not ic_df.empty else math.nan
        ic_mean = float(ic_df["rank_ic"].mean()) if not ic_df.empty else math.nan
        summary = {
            "source_backtest": str(source),
            "period_start": str(pd.to_datetime(eq_df["rebalance_date"].min()).date()) if not eq_df.empty else None,
            "period_end": str(pd.to_datetime(eq_df["rebalance_date"].max()).date()) if not eq_df.empty else None,
            "periods": int(len(eq_df)),
            "stocks_per_period_avg": float(ic_df["stocks"].mean()) if not ic_df.empty else None,
            "rank_ic": {
                "mean": ic_mean,
                "std": ic_std,
                "icir": float(ic_mean / ic_std) if ic_std and not math.isnan(ic_std) else None,
                "annualized_icir_assuming_5d_rebalance": float(ic_mean / ic_std * math.sqrt(252 / 5))
                if ic_std and not math.isnan(ic_std)
                else None,
                "positive_rate": float((ic_df["rank_ic"] > 0).mean()) if not ic_df.empty else None,
                "min": float(ic_df["rank_ic"].min()) if not ic_df.empty else None,
                "max": float(ic_df["rank_ic"].max()) if not ic_df.empty else None,
            },
            "benchmark": {
                "strategy_top5_total_return": float(eq_df["strategy_equity"].iloc[-1] - 1.0) if not eq_df.empty else None,
                "equal_weight_total_return": float(eq_df["equal_weight_equity"].iloc[-1] - 1.0) if not eq_df.empty else None,
                "strategy_minus_equal_weight_equity_ratio_total": float(eq_df["excess_equity_ratio"].iloc[-1] - 1.0)
                if not eq_df.empty
                else None,
                "avg_period_excess_return": float(eq_df["excess_return"].mean()) if not eq_df.empty else None,
                "excess_win_rate": float((eq_df["excess_return"] > 0).mean()) if not eq_df.empty else None,
            },
            "q5_stats": q5_stats.to_dict("records"),
            "q10_stats": q10_stats.to_dict("records"),
        }
        write_json(out / "diagnostics_summary.json", summary)


def score_feature_groups(run_dir: Path) -> None:
    for group in TOPK_GROUPS:
        feature_dir = run_dir / "feature_group_tests" / group
        cols = pd.read_parquet(feature_dir / "features.parquet").columns.tolist()
        keep = [c for c in cols if c not in {"future_return", "label"}]
        inf = pd.read_parquet(run_dir / "inference_features_latest.parquet")
        inf = inf[[c for c in keep if c in inf.columns]]
        inf_path = feature_dir / "inference_features_latest.parquet"
        inf.to_parquet(inf_path, index=False)
        run_py(
            [
                "train_lightgbm.py",
                "--train-path",
                str(feature_dir / "features.parquet"),
                "--inference-path",
                str(inf_path),
                "--model-dir",
                str(feature_dir / "models"),
                "--valid-days",
                "60",
                "--threshold",
                "0.5",
                "--top-k",
                "20",
            ]
        )


def add_df(doc: Any, df: pd.DataFrame, max_rows: int = 12) -> None:
    table_df = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(table_df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(table_df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            cells[i].text = str(value)


def latest_top(run_dir: Path, group: str = "momentum_liquidity", n: int = 20) -> pd.DataFrame:
    path = run_dir / "feature_group_tests" / group / "models" / "inference_scores_latest.parquet"
    if not path.exists():
        path = run_dir / "models" / "inference_scores_latest.parquet"
    df = pd.read_parquet(path).head(n)
    cols = [c for c in ["date", "code", "name", "industry", "close", "pct_chg_5d", "turnover", "score"] if c in df.columns]
    return df[cols].copy()


def generate_report(run_dir: Path, *, universe_label: str) -> Path:
    from docx import Document

    report_dir = run_dir.parent / "comparison_reports"
    report_dir.mkdir(parents=True, exist_ok=True)
    out_path = report_dir / f"{run_dir.name}_trader_report.docx"

    run_meta = read_json(run_dir / "run_meta.json")
    feat_meta = read_json(run_dir / "ml_features_ready.parquet.meta.json")
    train_meta = read_json(run_dir / "models" / "training_metadata.json")
    base = read_json(run_dir / "backtest" / "summary.json")
    group_results = read_json(run_dir / "feature_group_tests" / "group_run_results.json")
    topk = pd.read_csv(run_dir / "topk_tests" / "topk_results.csv")
    diag_ml = read_json(run_dir / "diagnostics" / "momentum_liquidity" / "diagnostics_summary.json")
    diag_vm = read_json(run_dir / "diagnostics" / "valuation_momentum" / "diagnostics_summary.json")

    group_table = (
        pd.DataFrame(
            [
                {
                    "特征组": k,
                    "总收益": pct(v.get("summary", {}).get("portfolio_total_return")),
                    "CAGR": pct(v.get("summary", {}).get("portfolio_cagr")),
                    "最大回撤": pct(v.get("summary", {}).get("portfolio_max_drawdown")),
                    "胜率": pct(v.get("summary", {}).get("portfolio_win_rate")),
                    "AUC": f3(v.get("summary", {}).get("oos_metrics", {}).get("auc")),
                }
                for k, v in group_results.items()
                if v.get("ok")
            ]
        )
        .assign(_sort=lambda x: x["总收益"].str.replace("%", "").astype(float))
        .sort_values("_sort", ascending=False)
        .drop(columns="_sort")
    )

    topk_table = topk[topk["ok"] == True].copy()  # noqa: E712
    topk_table["总收益"] = topk_table["total_return"].map(pct)
    topk_table["最大回撤"] = topk_table["max_drawdown"].map(pct)
    topk_table["胜率"] = topk_table["win_rate"].map(pct)
    topk_table = topk_table.sort_values("total_return", ascending=False)[["group", "top_k", "总收益", "最大回撤", "胜率"]]

    doc = Document()
    doc.add_heading(f"{universe_label} 两年机器学习选股全流程分析报告", level=0)
    doc.add_paragraph("第一页：Executive Summary")
    doc.add_heading("Executive Summary", level=1)
    best = topk[topk["ok"] == True].sort_values("total_return", ascending=False).iloc[0]  # noqa: E712
    doc.add_paragraph(
        f"交易员一句话：{universe_label}这次更像一个高弹性小盘/微盘股票池实验，"
        f"最佳组合是 {best['group']} Top{int(best['top_k'])}，样本期总收益 {pct(best['total_return'])}，"
        f"最大回撤 {pct(best['max_drawdown'])}，但执行层面必须把流动性、ST、涨跌停和换手成本放到第一优先级。"
    )
    doc.add_paragraph(
        f"基础全特征 Top5 总收益 {pct(base['portfolio_total_return'])}，CAGR {pct(base['portfolio_cagr'])}，"
        f"最大回撤 {pct(base['portfolio_max_drawdown'])}，调仓胜率 {pct(base['portfolio_win_rate'])}。"
    )
    doc.add_paragraph(
        f"机器学习视角：验证 AUC 为 {f3(train_meta['metrics']['auc'])}，walk-forward OOS AUC 为 "
        f"{f3(base['oos_metrics']['auc'])}。AUC不是本报告重点，重点是分层、TopK稳定性和交易可落地性。"
    )
    doc.add_paragraph(
        f"信号诊断：momentum_liquidity Rank IC 均值 {f3(diag_ml['rank_ic']['mean'])}，正 IC 比例 "
        f"{pct(diag_ml['rank_ic']['positive_rate'])}；valuation_momentum Rank IC 均值 "
        f"{f3(diag_vm['rank_ic']['mean'])}。如果最高分组没有持续强于低分组，交易上就不能只看TopK收益。"
    )
    doc.add_paragraph(
        "交易结论：这套结果适合作为候选池和盘前观察清单，不适合直接裸跑实盘。中证2000天然更容易出现小成交、ST、连板和跌停出不来的问题，下一版应先加入交易成本、停牌/ST过滤、涨跌停可成交性和容量约束。"
    )

    doc.add_heading("1. 实验口径", level=1)
    add_df(
        doc,
        pd.DataFrame(
            [
                ["指数", universe_label],
                ["指数代码", run_meta["index_code"]],
                ["成分日期", run_meta["constituent_date_max"]],
                ["样本窗口", f"{run_meta['start_date']} - {run_meta['end_date']}"],
                ["成分股数量", run_meta["stock_count"]],
                ["下载成功", run_meta["download_success_count"]],
                ["训练样本", feat_meta["train_rows"]],
                ["训练股票数", feat_meta["train_codes"]],
                ["回测区间", f"{base['backtest_start']} - {base['backtest_end']}"],
                ["调仓频率", "5个交易日"],
            ],
            columns=["项目", "值"],
        ),
    )

    doc.add_heading("2. 交易员视角核心表现", level=1)
    add_df(doc, group_table)
    doc.add_paragraph(
        "看这张表时，优先顺序不是AUC，而是：总收益是否来自可解释的风格、回撤是否可忍受、TopK扩容后是否仍然有效。"
    )

    doc.add_heading("3. TopK 容量测试", level=1)
    add_df(doc, topk_table, max_rows=21)
    doc.add_paragraph(
        "Top1如果特别强但回撤也极大，通常更像单票运气；Top5/Top10/Top15同时较强，才更像可组合化交易信号。"
    )

    doc.add_heading("4. 分层与 Rank IC", level=1)
    add_df(
        doc,
        pd.DataFrame(
            [
                ["momentum_liquidity", f3(diag_ml["rank_ic"]["mean"]), f3(diag_ml["rank_ic"]["icir"]), pct(diag_ml["rank_ic"]["positive_rate"]), pct(diag_ml["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"])],
                ["valuation_momentum", f3(diag_vm["rank_ic"]["mean"]), f3(diag_vm["rank_ic"]["icir"]), pct(diag_vm["rank_ic"]["positive_rate"]), pct(diag_vm["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"])],
            ],
            columns=["信号组", "Rank IC均值", "ICIR", "正IC比例", "相对等权权益超额"],
        ),
    )

    doc.add_heading("5. 最新观察名单", level=1)
    add_df(doc, latest_top(run_dir), max_rows=20)
    doc.add_paragraph("这不是买入指令，而是下一交易日盘前观察池。交易员需要再看停牌、ST、涨跌停距离、成交额和盘口冲击。")

    doc.add_heading("6. 风险与下一步", level=1)
    for text in [
        "中证2000的收益弹性高，但滑点和冲击成本也更高，Top5组合尤其容易被交易成本吃掉。",
        "指数全成分包含ST或准ST股票，报告按全成分研究保留，实盘应设置可交易过滤。",
        "当前回测用未来5日收益近似调仓收益，没有逐笔成交、涨跌停无法成交、停牌、手续费、印花税和滑点。",
        "机器学习部分有排序价值，但不是重点；如果交易分层不稳，AUC再好也不能直接落地。",
        "建议下一版做成本敏感性：单边10/20/30/50bp，Top5/10/15分别测容量，并加入成交额下限。",
    ]:
        doc.add_paragraph(text, style=None)

    doc.save(out_path)
    return out_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run index research pipeline.")
    parser.add_argument("--universe", default="csi2000")
    parser.add_argument("--index-code", default="932000")
    parser.add_argument("--label", default="中证2000")
    parser.add_argument("--run-dir", default="quant_data/csi2000_2y_run")
    parser.add_argument("--start-date", default="20240430")
    parser.add_argument("--end-date", default="20260430")
    parser.add_argument("--sleep", type=float, default=0.05)
    parser.add_argument("--overwrite", action="store_true")
    parser.add_argument("--skip-download", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    run_dir = ROOT / args.run_dir
    label_horizon = 5
    label_threshold = 0.02
    profile = f"{args.universe}_2y_5d"

    if not args.skip_download:
        download_index_data(
            run_dir=run_dir,
            universe=args.universe,
            index_code=args.index_code,
            start_date=args.start_date,
            end_date=args.end_date,
            sleep=args.sleep,
            overwrite=args.overwrite,
        )

    run_py(
        [
            "feature_engineering.py",
            "--data-dir",
            str(run_dir),
            "--output",
            str(run_dir / "ml_features_ready.parquet"),
            "--limit",
            "0",
            "--label-threshold",
            str(label_threshold),
            "--label-horizon",
            str(label_horizon),
            "--profile-name",
            profile,
        ]
    )
    run_py(
        [
            "build_inference_features.py",
            "--data-dir",
            str(run_dir),
            "--output",
            str(run_dir / "inference_features_latest.parquet"),
            "--limit",
            "0",
            "--as-of-date",
            "2026-04-30",
        ]
    )
    run_py(
        [
            "train_lightgbm.py",
            "--train-path",
            str(run_dir / "ml_features_ready.parquet"),
            "--inference-path",
            str(run_dir / "inference_features_latest.parquet"),
            "--model-dir",
            str(run_dir / "models"),
            "--valid-days",
            "60",
            "--threshold",
            "0.5",
            "--top-k",
            "20",
        ]
    )
    run_py(
        [
            "backtest_walk_forward.py",
            "--train-path",
            str(run_dir / "ml_features_ready.parquet"),
            "--output-dir",
            str(run_dir / "backtest"),
            "--min-train-days",
            "252",
            "--retrain-every",
            "20",
            "--rebalance-every",
            "5",
            "--top-k",
            "5",
            "--threshold",
            "0.5",
            "--profile-name",
            profile,
            "--profile-label",
            args.label,
            "--label-horizon",
            str(label_horizon),
            "--label-threshold",
            str(label_threshold),
        ]
    )

    make_feature_group_files(run_dir)
    run_feature_group_backtests(run_dir, label_horizon=label_horizon, label_threshold=label_threshold)
    run_topk_tests(run_dir, label_horizon=label_horizon, label_threshold=label_threshold)
    run_diagnostics(run_dir)
    score_feature_groups(run_dir)
    report_path = generate_report(run_dir, universe_label=args.label)

    print(f"REPORT: {report_path}", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
