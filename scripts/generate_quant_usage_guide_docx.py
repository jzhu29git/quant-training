#!/usr/bin/env python3
"""Generate a layperson-friendly Chinese usage guide for the quant system."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "quant_data" / "comparison_reports"
OUT_PATH = OUT_DIR / "quant_trade_cn_usage_guide.docx"

FONT_BODY = "Microsoft YaHei"
FONT_TITLE = "Microsoft YaHei UI"


def set_run_font(run, font_name: str = FONT_BODY, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def set_style_font(style, font_name: str, size: float | None = None) -> None:
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(run, FONT_TITLE, 18, True, (31, 78, 121))
    else:
        set_run_font(run, FONT_TITLE, 13 if level == 1 else 11.5, True, (31, 78, 121))


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, 10.5)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, "Consolas", 9.5)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        p = table.rows[0].cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, FONT_BODY, 9.5, True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, FONT_BODY, 9)


def build_doc() -> Document:
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"]:
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], FONT_BODY if style_name == "Normal" else FONT_TITLE, 10.5)

    add_heading(doc, "quant-trading-cn 大白话使用说明", level=0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("从模型原理到每5天怎么调仓，再到 Futu 模拟交易")
    set_run_font(run, FONT_BODY, 11, False, (88, 88, 88))

    add_heading(doc, "一句话", level=1)
    add_para(doc, "这套系统不是预测明天涨停，而是每隔5个交易日，把股票池里的所有股票重新打分，选出模型认为未来5个交易日更有机会跑赢的 TopK，然后按组合方式调仓。")
    add_para(doc, "它更像一个盘前选股和调仓助手，不是无人值守的实盘印钞机。")

    add_heading(doc, "系统流水线", level=1)
    add_table(
        doc,
        ["步骤", "做什么", "交易员怎么理解"],
        [
            ["Step 1", "下载日线、成交量、换手、估值数据", "准备行情和基础面板"],
            ["Step 2", "生成训练特征和未来5日标签", "告诉模型以前什么形态容易涨"],
            ["Step 3", "生成最新一天的推理特征", "准备今天要打分的股票表"],
            ["Step 4", "训练 LightGBM 并给最新截面打分", "生成最新 Top10/Top20 候选"],
            ["Step 5", "walk-forward 回测", "模拟历史上每5天调仓一次的效果"],
            ["Step 6", "Futu paper trading", "把最新候选转成模拟目标持仓和订单"],
        ],
    )

    add_heading(doc, "模型到底学什么", level=1)
    add_para(doc, "训练标签很简单：如果今天买入一只股票，未来5个交易日收益是否超过2%。超过就是1，否则就是0。")
    add_para(doc, "所以模型学的是：这只股票未来5天涨超2%的概率高不高。它不是直接预测明天涨跌，也不是预测具体涨幅。")
    add_para(doc, "LightGBM 可以理解成很多小交易规则树的集合。它会学习类似这样的组合条件：短期跌幅、20日趋势、换手、成交额、波动率、估值等一起出现时，未来5日更容易上涨还是不容易上涨。")

    add_heading(doc, "因子和权重会不会动态变化", level=1)
    add_para(doc, "因子集合不会自动变化。比如 momentum_liquidity 就固定使用均线、涨跌幅、成交额、换手等一组特征。模型不会自己突然发明 RSI、MACD 或行业强弱，除非我们写代码加进去。")
    add_para(doc, "模型内部的规则和非线性权重会随着重新训练而变化。同一个因子，这次训练可能很重要，下次训练可能不那么重要。")
    add_para(doc, "超参数也基本固定，例如5日标签、涨超2%作为正样本、TopK、树模型参数等，不会在回测中自动优化。")
    add_para(doc, "LightGBM 没有一个简单的线性权重表，不是 30%动量 + 20%换手 + 10%估值。我们能看的主要是 feature_importance，也就是特征重要性。")

    add_heading(doc, "OOS walk-forward 是什么", level=1)
    add_para(doc, "OOS walk-forward 的意思是：站在历史上的某一天，只用那天以前的数据训练模型，然后预测那天股票的排名。过5个交易日后，再看这些股票真实表现。")
    add_para(doc, "它的目的就是避免偷看未来。回测里 future_return 只用于事后验证，不用于当天选股。")
    add_para(doc, "当前逻辑是每5个交易日调仓一次，每20次调仓重新训练一次模型，大约相当于每100个交易日重新训练一次。")

    add_heading(doc, "5天后具体怎么操作", level=1)
    add_para(doc, "假设今天模型给出 Top10，标准做法是买入可交易的 Top10，并持有5个交易日。")
    add_para(doc, "5个交易日后，重新更新数据并重新生成 Top10。然后做再平衡：仍在新Top10里的继续持有，不在新Top10里的卖出，新进入Top10的买入。")
    add_para(doc, "如果做等权 Top10，每只目标仓位约10%；如果做 Top5，每只约20%。中证2000不建议Top5过度集中，中证500 Top5/Top10更值得继续研究。")
    add_code(doc, "旧Top10 ∩ 新Top10：继续持有或微调仓位")
    add_code(doc, "旧Top10 不在 新Top10：卖出")
    add_code(doc, "新Top10 但当前没有：买入")

    add_heading(doc, "日常使用建议", level=1)
    add_para(doc, "先不要一上来全自动实盘。建议先做人工辅助交易：每5个交易日更新数据、生成Top10、人工过滤不可交易票，再手动或模拟下单。")
    add_para(doc, "人工过滤至少包括：ST、停牌、接近涨停、接近跌停、成交额太小、一字板、财报雷、盘口太薄。")
    add_para(doc, "等模拟跑顺以后，再考虑让 Futu paper trading 自动同步目标仓位。")

    add_heading(doc, "常用命令示例", level=1)
    add_para(doc, "如果继续使用我们新做的指数研究脚本，可以这样跑中证2000全流程：")
    add_code(doc, "python scripts\\run_index_research_pipeline.py --universe csi2000 --index-code 932000 --label 中证2000 --run-dir quant_data/csi2000_2y_run --start-date 20240430 --end-date 20260430 --sleep 0.05")
    add_para(doc, "如果只想重新生成交易员视角 Word 报告：")
    add_code(doc, "python scripts\\generate_three_index_trader_insights_docx.py")

    add_heading(doc, "Futu 模拟交易是什么", level=1)
    add_para(doc, "Futu 模拟交易不是模型本身，而是执行层。它读取最新 inference_scores_latest.parquet，选出分数超过阈值的 TopK，然后和当前持仓对比，生成买卖计划。")
    add_para(doc, "默认参数大致是：TopK=5，min_score=0.5，A股一手100股，保留2%现金，买卖限价各给50bp空间，单笔最大1000股。")
    add_para(doc, "daemon 是后台循环，每隔一段时间检查最新分数文件是否变化。分数文件没变就不重复交易；变了就重新同步目标仓位。")
    add_para(doc, "注意：paper_trade_futu.py 会调用外部 gateway 的下单接口。是否真钱取决于 gateway 后面接的是模拟账户还是真实账户。没有确认前必须先用 dry-run 或只看 targets_latest.parquet。")

    add_heading(doc, "当前系统最大的不足", level=1)
    for text in [
        "当前回测没有充分计入真实交易成本。",
        "没有严格模拟涨跌停无法成交。",
        "中证2000还需要先过滤 ST、停牌和极低流动性股票。",
        "没有行业暴露控制。",
        "没有完整止损/止盈规则。",
        "目前更像选股器和组合研究工具，还不是完整无人值守交易系统。",
    ]:
        add_para(doc, text)

    add_heading(doc, "推荐下一步", level=1)
    add_table(
        doc,
        ["顺序", "任务", "目的"],
        [
            ["1", "固定主策略为中证500 momentum_liquidity Top10", "先把最像可交易的方向打磨好"],
            ["2", "加入 ST、停牌、涨跌停过滤", "避免纸面收益无法成交"],
            ["3", "加入成交额和容量过滤", "控制滑点和冲击成本"],
            ["4", "加入手续费、印花税、滑点回测", "看真实收益还剩多少"],
            ["5", "用 Futu dry-run 生成目标仓位", "先验证交易计划，不直接下单"],
            ["6", "跑一段模拟交易", "观察实际换手、成交和回撤"],
        ],
    )

    return doc


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_PATH)
    print(OUT_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
