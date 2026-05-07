from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "quant_data" / "comparison_reports"
OUT_PATH = OUT_DIR / "sse50_vs_csi500_2y_report_clean.docx"


def read_json(rel_path: str) -> dict:
    return json.loads((ROOT / rel_path).read_text(encoding="utf-8"))


def pct(value: float) -> str:
    return f"{float(value) * 100:.1f}%"


def f3(value: float) -> str:
    return f"{float(value):.3f}"


def feature_group_table(run_name: str) -> pd.DataFrame:
    payload = read_json(f"quant_data/{run_name}/feature_group_tests/group_run_results.json")
    rows = []
    for name, item in payload.items():
        if not item.get("ok"):
            continue
        summary = item["summary"]
        rows.append(
            {
                "特征组": name,
                "总收益": pct(summary["portfolio_total_return"]),
                "年化": pct(summary["portfolio_cagr"]),
                "最大回撤": pct(summary["portfolio_max_drawdown"]),
                "胜率": pct(summary["portfolio_win_rate"]),
                "OOS AUC": f3(summary["oos_metrics"]["auc"]),
                "_sort": summary["portfolio_total_return"],
            }
        )
    return pd.DataFrame(rows).sort_values("_sort", ascending=False).drop(columns=["_sort"])


def topk_table(run_name: str, n: int = 10) -> pd.DataFrame:
    df = pd.read_csv(ROOT / f"quant_data/{run_name}/topk_tests/topk_results.csv")
    df = df[df["ok"] == True].copy()
    df["收益回撤比"] = df["total_return"] / df["max_drawdown"].abs()
    df = df.sort_values("total_return", ascending=False).head(n)
    return pd.DataFrame(
        {
            "特征组": df["group"],
            "TopK": df["top_k"].astype(int),
            "总收益": df["total_return"].map(pct),
            "年化": df["cagr"].map(pct),
            "最大回撤": df["max_drawdown"].map(pct),
            "胜率": df["win_rate"].map(pct),
            "收益回撤比": df["收益回撤比"].map(lambda x: f"{x:.2f}"),
        }
    )


def diagnostics(run_name: str, group: str) -> dict:
    return read_json(f"quant_data/{run_name}/diagnostics/{group}/diagnostics_summary.json")


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(20)


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Microsoft YaHei"
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
    doc.add_paragraph("")


def add_df(doc: Document, df: pd.DataFrame) -> None:
    add_table(doc, list(df.columns), df.astype(str).values.tolist())


def bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    sse_run = read_json("quant_data/sse50_2y_run/run_meta.json")
    csi_run = read_json("quant_data/csi500_2y_run/run_meta.json")
    sse_feat = read_json("quant_data/sse50_2y_run/ml_features_ready.parquet.meta.json")
    csi_feat = read_json("quant_data/csi500_2y_run/ml_features_ready.parquet.meta.json")
    sse_train = read_json("quant_data/sse50_2y_run/models/training_metadata.json")
    csi_train = read_json("quant_data/csi500_2y_run/models/training_metadata.json")
    sse_base = read_json("quant_data/sse50_2y_run/backtest/summary.json")
    csi_base = read_json("quant_data/csi500_2y_run/backtest/summary.json")
    sse_diag = diagnostics("sse50_2y_run", "valuation_momentum")
    csi_diag_ml = diagnostics("csi500_2y_run", "momentum_liquidity")
    csi_diag_vm = diagnostics("csi500_2y_run", "valuation_momentum")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        doc.styles[style_name].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    add_title(doc, "上证50 vs 中证500：两年样本 LightGBM 选股实验对比报告")
    doc.add_paragraph("生成日期：2026-04-30")
    doc.add_paragraph(
        "口径：当前成分股、2024-04-30 到 2026-04-30 日频数据、标签为未来 5 个交易日收益超过 2%。"
        "所有回测均为 walk-forward 样本外预测；没有交易成本；没有历史动态成分股。"
    )

    doc.add_heading("第 1 页：大白话 Executive Summary", level=1)
    doc.add_paragraph(
        "这次实验可以理解成：我们用同一套机器学习选股方法，分别在上证50和中证500两个股票池里试了一遍。"
        "上证50只有50只大盘股，股票之间差异相对小；中证500有500只中盘股，横截面更宽，模型更容易找到相对强弱。"
    )
    bullet(
        doc,
        [
            f"上证50：训练样本 {sse_feat['train_rows']:,} 行，验证 AUC {sse_train['metrics']['auc']:.3f}。最值得关注的因子组是 valuation_momentum。",
            f"中证500：训练样本 {csi_feat['train_rows']:,} 行，验证 AUC {csi_train['metrics']['auc']:.3f}。最值得关注的因子组变成 momentum_liquidity。",
            f"上证50 valuation_momentum Top5 总收益 {pct(sse_diag['benchmark']['strategy_top5_total_return'])}，等权基准 {pct(sse_diag['benchmark']['equal_weight_total_return'])}。",
            f"中证500 momentum_liquidity Top5 总收益 {pct(csi_diag_ml['benchmark']['strategy_top5_total_return'])}，等权基准 {pct(csi_diag_ml['benchmark']['equal_weight_total_return'])}。",
            "最大的新发现：liquidity 在上证50里贡献不明显，但在中证500里非常重要。",
            "最大风险：两个实验都用了当前成分股回看历史，有幸存者偏差；没有交易成本；只有约42个调仓期。",
        ],
    )
    doc.add_paragraph(
        "一句话判断：中证500比上证50更适合这套 LightGBM 横截面选股框架；"
        "但下一步必须加入交易成本、历史动态成分股、真正独立测试集。"
    )

    doc.add_heading("第 2 页：实验设置", level=1)
    bullet(
        doc,
        [
            "股票池：上证50指数 000016，中证500指数 000905。",
            "成分股来源：AkShare 中证指数成分股接口。",
            "行情与估值来源：BaoStock。",
            "时间区间：2024-04-30 到 2026-04-30。",
            "标签定义：未来5个交易日收益超过2%记为正样本。",
            "调仓频率：每5个交易日调仓一次。",
            "回测方法：至少252个交易日训练，然后 walk-forward 样本外预测。",
        ],
    )

    doc.add_heading("第 3 页：运行时间和样本规模", level=1)
    add_table(
        doc,
        ["项目", "上证50", "中证500"],
        [
            ["股票数量", sse_run["stock_count"], csi_run["stock_count"]],
            ["下载成功", sse_run["download_success_count"], csi_run["download_success_count"]],
            ["下载耗时", f"{sse_run['elapsed_seconds']:.1f} 秒", f"{csi_run['elapsed_seconds']:.1f} 秒"],
            ["训练样本", f"{sse_feat['train_rows']:,}", f"{csi_feat['train_rows']:,}"],
            ["可训练日期", f"{sse_feat['date_min']} 至 {sse_feat['date_max']}", f"{csi_feat['date_min']} 至 {csi_feat['date_max']}"],
        ],
    )
    doc.add_paragraph("中证500下载耗时约为上证50的10.7倍，基本和股票数量比例一致；但训练和回测仍然是秒到分钟级。")

    doc.add_heading("第 4 页：训练验证指标", level=1)
    add_table(
        doc,
        ["指标", "上证50全特征", "中证500全特征"],
        [
            [key, f3(sse_train["metrics"][key]), f3(csi_train["metrics"][key])]
            for key in ["auc", "accuracy", "precision", "recall", "positive_rate"]
        ],
    )
    doc.add_paragraph("中证500验证 AUC 从上证50的0.528提升到0.565。更大的横截面给模型更多相对排序样本。")

    doc.add_heading("第 5 页：全特征 baseline 回测", level=1)
    add_table(
        doc,
        ["指标", "上证50全特征Top5", "中证500全特征Top5"],
        [
            ["总收益", pct(sse_base["portfolio_total_return"]), pct(csi_base["portfolio_total_return"])],
            ["年化", pct(sse_base["portfolio_cagr"]), pct(csi_base["portfolio_cagr"])],
            ["最大回撤", pct(sse_base["portfolio_max_drawdown"]), pct(csi_base["portfolio_max_drawdown"])],
            ["胜率", pct(sse_base["portfolio_win_rate"]), pct(csi_base["portfolio_win_rate"])],
            ["OOS AUC", f3(sse_base["oos_metrics"]["auc"]), f3(csi_base["oos_metrics"]["auc"])],
        ],
    )
    doc.add_paragraph("中证500 baseline 更强：收益更高、AUC更高，回撤和上证50接近。")

    doc.add_heading("第 6 页：特征组对照", level=1)
    doc.add_heading("上证50特征组 Top5 排名", level=2)
    add_df(doc, feature_group_table("sse50_2y_run"))
    doc.add_heading("中证500特征组 Top5 排名", level=2)
    add_df(doc, feature_group_table("csi500_2y_run"))
    doc.add_paragraph("上证50最强是 valuation_momentum，中证500最强是 momentum_liquidity。")

    doc.add_heading("第 7 页：TopK 敏感性", level=1)
    doc.add_heading("上证50 TopK 总收益前10", level=2)
    add_df(doc, topk_table("sse50_2y_run"))
    doc.add_heading("中证500 TopK 总收益前10", level=2)
    add_df(doc, topk_table("csi500_2y_run"))
    doc.add_paragraph("上证50里 TopK 越小越强；中证500里 Top1 不稳，Top5/10/15 更合理。")

    doc.add_heading("第 8 页：分层收益对比", level=1)
    doc.add_heading("上证50 valuation_momentum 五分组", level=2)
    add_df(doc, pd.DataFrame(sse_diag["q5_stats"]))
    doc.add_heading("中证500 momentum_liquidity 五分组", level=2)
    add_df(doc, pd.DataFrame(csi_diag_ml["q5_stats"]))
    doc.add_paragraph("中证500的五分组更健康：Q1到Q5大体递增；上证50是Q5最强但中间层较乱。")

    doc.add_heading("第 9 页：Rank IC / ICIR", level=1)
    add_table(
        doc,
        ["指标", "上证50 valuation_momentum", "中证500 momentum_liquidity", "中证500 valuation_momentum"],
        [
            ["平均 Rank IC", f3(sse_diag["rank_ic"]["mean"]), f3(csi_diag_ml["rank_ic"]["mean"]), f3(csi_diag_vm["rank_ic"]["mean"])],
            ["Rank IC 标准差", f3(sse_diag["rank_ic"]["std"]), f3(csi_diag_ml["rank_ic"]["std"]), f3(csi_diag_vm["rank_ic"]["std"])],
            ["ICIR", f3(sse_diag["rank_ic"]["icir"]), f3(csi_diag_ml["rank_ic"]["icir"]), f3(csi_diag_vm["rank_ic"]["icir"])],
            [
                "年化ICIR",
                f3(sse_diag["rank_ic"]["annualized_icir_assuming_5d_rebalance"]),
                f3(csi_diag_ml["rank_ic"]["annualized_icir_assuming_5d_rebalance"]),
                f3(csi_diag_vm["rank_ic"]["annualized_icir_assuming_5d_rebalance"]),
            ],
            ["Rank IC为正比例", pct(sse_diag["rank_ic"]["positive_rate"]), pct(csi_diag_ml["rank_ic"]["positive_rate"]), pct(csi_diag_vm["rank_ic"]["positive_rate"])],
        ],
    )

    doc.add_heading("第 10 页：等权 benchmark 与 alpha 迹象", level=1)
    add_table(
        doc,
        ["指标", "上证50 valuation_momentum Top5", "中证500 momentum_liquidity Top5", "中证500 valuation_momentum Top5"],
        [
            ["策略Top5总收益", pct(sse_diag["benchmark"]["strategy_top5_total_return"]), pct(csi_diag_ml["benchmark"]["strategy_top5_total_return"]), pct(csi_diag_vm["benchmark"]["strategy_top5_total_return"])],
            ["等权基准总收益", pct(sse_diag["benchmark"]["equal_weight_total_return"]), pct(csi_diag_ml["benchmark"]["equal_weight_total_return"]), pct(csi_diag_vm["benchmark"]["equal_weight_total_return"])],
            ["相对等权权益超额", pct(sse_diag["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"]), pct(csi_diag_ml["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"]), pct(csi_diag_vm["benchmark"]["strategy_minus_equal_weight_equity_ratio_total"])],
            ["平均每期超额", pct(sse_diag["benchmark"]["avg_period_excess_return"]), pct(csi_diag_ml["benchmark"]["avg_period_excess_return"]), pct(csi_diag_vm["benchmark"]["avg_period_excess_return"])],
            ["超额为正比例", pct(sse_diag["benchmark"]["excess_win_rate"]), pct(csi_diag_ml["benchmark"]["excess_win_rate"]), pct(csi_diag_vm["benchmark"]["excess_win_rate"])],
        ],
    )

    doc.add_heading("第 11 页：为什么中证500里 liquidity 变重要", level=1)
    doc.add_paragraph(
        "上证50股票普遍流动性强，换手和成交量差异没有那么能区分股票；中证500成分更分散，"
        "量能变化、换手变化可能代表资金关注度、短期拥挤度、趋势确认或流动性改善。"
    )

    doc.add_heading("第 12 页：为什么 AUC 不高但收益高", level=1)
    doc.add_paragraph(
        "AUC 看全样本排序，策略只买 TopK。模型可能无法把500只股票完整排得很顺，"
        "但只要能把最值得买的一小撮挑出来，组合收益就可能好。"
    )

    doc.add_heading("第 13 页：主要风险和偏差", level=1)
    bullet(
        doc,
        [
            "幸存者偏差：使用当前成分股回测历史，历史上被剔除的股票没有进入样本。",
            "没有交易成本：TopK越小、换手越高，成本影响越大。",
            "样本短：只有约42个调仓期。",
            "参数选择偏差：已经试过多个特征组和TopK，最佳结果可能被挑选偏差放大。",
            "没有行业/市值中性：收益可能包含风格和行业暴露。",
        ],
    )

    doc.add_heading("第 14 页：实盘意义和暂定方向", level=1)
    doc.add_paragraph(
        "如果只看研究优先级，我会把中证500放在上证50前面。中证500的可选股票更多，模型更容易挑出尾部差异；"
        "上证50成分太少，TopK=1/3看起来强但集中度过高。"
    )

    doc.add_heading("第 15 页：下一步建议", level=1)
    bullet(
        doc,
        [
            "加入交易成本：手续费、印花税、滑点、冲击成本。",
            "做历史动态成分股，消除当前成分股回看历史的幸存者偏差。",
            "做真正训练/验证/测试三段，不再用测试期挑参数。",
            "对中证500做行业、市值、波动率暴露拆解。",
            "做滚动年度验证，看2024、2025、2026分别是否有效。",
            "尝试 LightGBM Ranker，因为选股本质是排序，不是二分类。",
            "做换手率和容量分析，确认TopK=5是否能实盘承载。",
        ],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
