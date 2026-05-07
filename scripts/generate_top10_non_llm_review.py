#!/usr/bin/env python3
"""Generate a non-LLM trader review for latest TopK candidates."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

import pandas as pd
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CONFIG_PATH = ROOT / "paper_trading_config.yaml"
OUT_DIR = ROOT / "quant_data" / "comparison_reports"
FONT_BODY = "Microsoft YaHei"
FONT_TITLE = "Microsoft YaHei"


def read_config(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as handle:
        data = yaml.safe_load(handle) or {}
    if not isinstance(data, dict):
        raise ValueError(f"config must be a mapping: {path}")
    return data


def deep_get(payload: dict[str, Any], keys: list[str], default: Any = None) -> Any:
    current: Any = payload
    for key in keys:
        if not isinstance(current, dict) or key not in current:
            return default
        current = current[key]
    return current


def to_float(value: Any, default: float = 0.0) -> float:
    try:
        if value in (None, "", "N/A"):
            return default
        return float(value)
    except (TypeError, ValueError):
        return default


def pct(value: Any) -> str:
    try:
        return f"{float(value):.1f}%"
    except Exception:
        return "N/A"


def money_cn(value: Any) -> str:
    amount = to_float(value)
    if amount >= 100_000_000:
        return f"{amount / 100_000_000:.2f}亿"
    if amount >= 10_000:
        return f"{amount / 10_000:.0f}万"
    return f"{amount:.0f}"


def normalize_code(value: Any) -> str:
    text = str(value or "").strip()
    return text.zfill(6) if text.isdigit() else text


def load_latest_candidates(run_dir: Path, group: str, top_k: int) -> pd.DataFrame:
    score_path = run_dir / "feature_group_tests" / group / "models" / "inference_scores_latest.parquet"
    if not score_path.exists():
        score_path = run_dir / "models" / "inference_scores_latest.parquet"
    if not score_path.exists():
        raise FileNotFoundError(f"missing score file: {score_path}")

    scores = pd.read_parquet(score_path)
    scores["code"] = scores["code"].astype(str).str.zfill(6)
    scores["score"] = pd.to_numeric(scores["score"], errors="coerce")
    scores = scores.dropna(subset=["score"]).sort_values("score", ascending=False).head(top_k).copy()
    scores["rank"] = range(1, len(scores) + 1)

    feature_path = run_dir / "inference_features_latest.parquet"
    if feature_path.exists():
        features = pd.read_parquet(feature_path)
        features["code"] = features["code"].astype(str).str.zfill(6)
        keep_cols = [
            "code",
            "open",
            "high",
            "low",
            "close",
            "amount",
            "volume",
            "turnover",
            "pct_chg",
            "pct_chg_5d",
            "pct_chg_20d",
            "bias_20",
            "turnover_ma5",
            "volume_ma5",
            "volatility_20d",
            "close_to_high_20d",
            "close_to_low_20d",
            "pe_ttm",
            "pb",
            "ps",
            "pcf",
        ]
        features = features[[col for col in keep_cols if col in features.columns]].drop_duplicates("code", keep="last")
        scores = scores.merge(features, on="code", how="left", suffixes=("", "_feature"))
        for col in ["open", "high", "low", "close", "amount", "volume", "turnover", "pct_chg", "pct_chg_5d", "pct_chg_20d", "bias_20", "turnover_ma5", "volume_ma5"]:
            feature_col = f"{col}_feature"
            if feature_col in scores.columns:
                if col in scores.columns:
                    scores[col] = scores[col].combine_first(scores[feature_col])
                else:
                    scores[col] = scores[feature_col]
                scores = scores.drop(columns=[feature_col])
    return scores


def classify_liquidity(amount: float, turnover: float) -> tuple[str, list[str]]:
    flags: list[str] = []
    if amount >= 300_000_000:
        label = "强"
    elif amount >= 100_000_000:
        label = "可交易"
    elif amount >= 50_000_000:
        label = "偏薄"
        flags.append("成交额偏低")
    else:
        label = "薄"
        flags.append("流动性不足")
    if turnover > 10:
        flags.append("换手过热")
    elif turnover < 0.5:
        flags.append("换手偏低")
    return label, flags


def review_row(row: pd.Series) -> dict[str, Any]:
    name = str(row.get("name") or "")
    code = normalize_code(row.get("code"))
    score = to_float(row.get("score"))
    amount = to_float(row.get("amount"))
    turnover = to_float(row.get("turnover"))
    turnover_ma5 = to_float(row.get("turnover_ma5"))
    pct_1d = to_float(row.get("pct_chg"))
    pct_5d = to_float(row.get("pct_chg_5d")) * 100 if abs(to_float(row.get("pct_chg_5d"))) < 2 else to_float(row.get("pct_chg_5d"))
    pct_20d = to_float(row.get("pct_chg_20d")) * 100 if abs(to_float(row.get("pct_chg_20d"))) < 2 else to_float(row.get("pct_chg_20d"))
    bias_20 = to_float(row.get("bias_20")) * 100 if abs(to_float(row.get("bias_20"))) < 2 else to_float(row.get("bias_20"))
    close_to_high_20d = to_float(row.get("close_to_high_20d")) * 100
    close_to_low_20d = to_float(row.get("close_to_low_20d")) * 100

    liquidity_label, flags = classify_liquidity(amount, turnover)
    positives: list[str] = []

    if score >= 0.6:
        positives.append("模型分数强")
    elif score >= 0.53:
        positives.append("模型分数较好")

    if amount >= 300_000_000:
        positives.append("成交额充足")
    if turnover_ma5 > 0 and turnover > turnover_ma5 * 1.5:
        flags.append("当日换手显著放大")
    if pct_1d <= -9.5:
        flags.append("接近跌停/大阴线")
    elif pct_1d >= 9.5:
        flags.append("接近涨停，追价风险")
    if pct_5d <= -15:
        flags.append("5日跌幅较深，可能是下跌中继")
    elif pct_5d >= 18:
        flags.append("5日涨幅过大，短线过热")
    if pct_20d <= -25:
        flags.append("20日弱势明显")
    elif pct_20d >= 35:
        flags.append("20日涨幅过大")
    if bias_20 <= -15:
        positives.append("低于20日均线较多，有反弹博弈")
    elif bias_20 >= 15:
        flags.append("偏离20日均线较高")
    if close_to_high_20d > -2 and close_to_high_20d != 0:
        flags.append("接近20日高位")
    if close_to_low_20d < 3 and close_to_low_20d != 0:
        positives.append("接近20日低位")
    if "ST" in name.upper():
        flags.append("ST风险")

    hard_exclude = any(flag in flags for flag in ["ST风险", "流动性不足"])
    if hard_exclude:
        decision = "剔除"
    elif len(flags) >= 3:
        decision = "谨慎"
    else:
        decision = "保留"

    if decision == "保留":
        one_liner = "可进入paper候选，重点观察成交额和次日承接。"
    elif decision == "谨慎":
        one_liner = "信号有交易价值，但短线风险较多，建议降低权重或等待确认。"
    else:
        one_liner = "不建议进入本轮paper组合，除非人工确认风险可控。"

    return {
        "rank": int(row.get("rank") or 0),
        "code": code,
        "name": name,
        "score": round(score, 6),
        "close": round(to_float(row.get("close")), 4),
        "amount_cn": money_cn(amount),
        "turnover": round(turnover, 4),
        "pct_chg": round(pct_1d, 4),
        "pct_chg_5d_pct": round(pct_5d, 4),
        "pct_chg_20d_pct": round(pct_20d, 4),
        "bias_20_pct": round(bias_20, 4),
        "liquidity_label": liquidity_label,
        "positives": "；".join(positives) if positives else "模型排序靠前",
        "risk_flags": "；".join(flags) if flags else "暂无明显硬伤",
        "decision": decision,
        "trader_note": one_liner,
    }


def set_run_font(run, font_name: str = FONT_BODY, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def set_style_font(style, font_name: str, size: float | None = None) -> None:
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(run, FONT_TITLE, 18, True, (31, 78, 121))
    else:
        set_run_font(run, FONT_TITLE, 13 if level == 1 else 11.5, True, (31, 78, 121))


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, 10.5)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, FONT_BODY, 8.8, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            run = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(run, FONT_BODY, 8.2)


def build_doc(review: pd.DataFrame, *, index_name: str, group: str, output_csv: Path) -> Document:
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Title"]:
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], FONT_BODY, 10.5)

    add_heading(doc, f"{index_name.upper()} Top10 无LLM交易员复核", level=0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_text = str(review["date"].iloc[0])[:10] if "date" in review.columns and not review.empty else ""
    run = subtitle.add_run(f"因子组：{group} | 信号日期：{date_text}")
    set_run_font(run, FONT_BODY, 10.5, False, (88, 88, 88))

    add_heading(doc, "Executive Summary", level=1)
    counts = review["decision"].value_counts().to_dict()
    add_para(
        doc,
        f"本报告不调用LLM，只基于本地行情、动量、流动性和位置指标做规则化复核。"
        f"本轮Top10中，保留 {counts.get('保留', 0)} 只，谨慎 {counts.get('谨慎', 0)} 只，剔除 {counts.get('剔除', 0)} 只。",
    )
    add_para(doc, "这份报告适合放在paper trading dry-run之前，帮助我们识别过热、流动性不足、接近涨跌停、短线下跌中继等交易风险。")

    add_heading(doc, "复核规则", level=1)
    add_para(doc, "硬剔除：ST风险、成交额过低等。谨慎：多项短线风险同时出现，例如接近涨跌停、5日跌幅过深、20日弱势、换手过热。保留：模型排序靠前且没有明显硬伤。")

    add_heading(doc, "Top10复核表", level=1)
    rows = []
    for row in review.itertuples(index=False):
        rows.append(
            [
                row.rank,
                row.code,
                row.name,
                f"{row.score:.3f}",
                row.amount_cn,
                pct(row.pct_chg_5d_pct),
                pct(row.pct_chg_20d_pct),
                pct(row.bias_20_pct),
                row.liquidity_label,
                row.decision,
                row.risk_flags,
            ]
        )
    add_table(
        doc,
        ["Rank", "代码", "名称", "Score", "成交额", "5日", "20日", "Bias20", "流动性", "建议", "风险标签"],
        rows,
    )

    add_heading(doc, "逐只点评", level=1)
    for row in review.itertuples(index=False):
        add_heading(doc, f"{row.rank}. {row.code} {row.name} | {row.decision}", level=2)
        add_para(doc, f"优势：{row.positives}")
        add_para(doc, f"风险：{row.risk_flags}")
        add_para(doc, f"交易员备注：{row.trader_note}")

    add_heading(doc, "输出文件", level=1)
    add_para(doc, f"CSV复核表：{output_csv}")
    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate non-LLM TopK trader review.")
    parser.add_argument("--config", default=str(DEFAULT_CONFIG_PATH))
    parser.add_argument("--index", choices=["csi500", "csi2000", "sse50"], default="csi500")
    parser.add_argument("--group", default="")
    parser.add_argument("--top-k", type=int, default=0)
    parser.add_argument("--output-dir", default=str(OUT_DIR))
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    config = read_config(Path(args.config))
    index_cfg = deep_get(config, ["indices", args.index], {})
    run_dir = ROOT / str(index_cfg.get("run_dir"))
    group = args.group or str(index_cfg.get("group") or "momentum_liquidity")
    top_k = args.top_k or int(index_cfg.get("top_k") or 10)
    candidates = load_latest_candidates(run_dir, group, top_k)

    review_rows = []
    for _, row in candidates.iterrows():
        item = review_row(row)
        item["date"] = str(pd.to_datetime(row.get("date")).date()) if row.get("date") is not None else ""
        item["group"] = group
        item["index"] = args.index
        review_rows.append(item)
    review = pd.DataFrame(review_rows)

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    suffix = f"{args.index}_{group}_top{top_k}_non_llm_review"
    csv_path = out_dir / f"{suffix}.csv"
    docx_path = out_dir / f"{suffix}.docx"
    review.to_csv(csv_path, index=False, encoding="utf-8-sig")
    doc = build_doc(review, index_name=args.index, group=group, output_csv=csv_path)
    doc.save(docx_path)

    print(json.dumps({"csv": str(csv_path), "docx": str(docx_path), "rows": len(review)}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
