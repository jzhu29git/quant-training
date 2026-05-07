#!/usr/bin/env python3
"""Generate the three-index trader-insight Word report with safe UTF-8 text."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "quant_data" / "comparison_reports"
OUT_PATH = OUT_DIR / "three_index_trader_insights_corrected.docx"

FONT_BODY = "Microsoft YaHei"
FONT_TITLE = "Microsoft YaHei UI"


def set_run_font(run, font_name: str = FONT_BODY, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def set_style_font(style, font_name: str, size: float | None = None, bold: bool | None = None) -> None:
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold

    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 0:
        set_run_font(run, FONT_TITLE, 18, True, (31, 78, 121))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_run_font(run, FONT_TITLE, 14, True, (31, 78, 121))
    else:
        set_run_font(run, FONT_TITLE, 12, True, (31, 78, 121))
    return paragraph


def add_para(doc: Document, text: str, size: float = 10.5):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, size)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, FONT_BODY, 9.5, True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, FONT_BODY, 9)


def build_doc() -> Document:
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle"]:
        if style_name in doc.styles:
            font = FONT_BODY if style_name == "Normal" else FONT_TITLE
            set_style_font(doc.styles[style_name], font, 10.5)

    add_heading(doc, "三个指数两年选股实验：交易员视角 Insights", level=0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("上证50 / 中证500 / 中证2000 对照复盘，少讲模型，多讲交易含义")
    set_run_font(subtitle_run, FONT_BODY, 11, False, (88, 88, 88))

    add_heading(doc, "Executive Summary", level=1)
    for text in [
        "中证500是目前最值得继续推进的池子。",
        "中证2000有更宽的选择面，但可交易性折损会很大。",
        "上证50能做，但更像少数票的择时/风格交易，不像稳定横截面选股池。",
    ]:
        add_para(doc, text)

    add_table(
        doc,
        ["指数", "最像什么交易", "最优观察", "主要问题"],
        [
            ["上证50", "大票风格轮动 / 核心资产择时", "valuation_momentum Top1/3 很强", "股票太少，Top1味道太重，组合化空间有限"],
            ["中证500", "最像可组合化 alpha 池", "momentum_liquidity Top5 最强", "需要做成本和换手检验"],
            ["中证2000", "小票弹性池 / 高波动反弹池", "momentum_liquidity Top10/15/20 更合理", "ST、流动性、涨跌停、滑点会吃掉很多纸面收益"],
        ],
    )

    sections = [
        (
            "核心洞察 1：中证500 是三者里交易形态最好的",
            [
                "中证500的表现最舒服，不只是因为收益最高，而是因为它的交易结构更像真实可做的组合。",
                "中证500里最强的是 momentum_liquidity Top5：总收益约 147.3%，最大回撤约 -11.0%，胜率约 61.9%，相对等权超额很明显。",
                "这说明中盘股里资金行为有痕迹。量能、换手、短期动量这些东西能把股票分出层次。交易员看这个会比较有感觉：有资金进、有趋势、有分歧、有弹性，最后能形成组合收益。",
                "中证500还有一个优点：Top5 强，Top10/15/20 也不差。实盘不是永远只买5只，如果考虑容量、停牌、冲击成本、风控替换，Top10/15还能撑住，说明这个池子不是靠一两只妖股撑起来的。",
                "我的判断：中证500是当前最应该继续打磨的主战场。",
            ],
        ),
        (
            "核心洞察 2：中证2000不是没机会，而是纸面收益要大幅打折",
            [
                "中证2000看起来也不错，尤其 momentum_liquidity Top10：总收益约 81.1%，最大回撤约 -9.6%，胜率约 61.9%。",
                "但这里要比中证500谨慎很多。中证2000的收益更像小票弹性和急跌反弹共同贡献。最新Top10里，很多股票过去5日跌幅很大，模型明显偏向“高换手 + 急跌后可能修复”的结构。",
                "这类票交易上有机会，但也最容易出问题：跌停买不到、卖不掉；小成交票滑点很大；ST/准ST风险高；高开低走和冲高回落很多；盘口厚度不够，回测收益容易虚胖。",
                "本次中证2000还有数据口径限制：2000只成分里，BaoStock成功覆盖1969只，31只920xxx代码返回空数据。",
                "所以中证2000更像一个高弹性的战术池，不是首选主组合池。",
                "如果做它，我会倾向不做Top5，而重点看Top10/15/20；必须加成交额下限、过滤ST、过滤接近涨跌停、模拟滑点和不能成交，并使用更短持仓和更严格止损。",
            ],
        ),
        (
            "核心洞察 3：上证50不是没效果，但它不是横截面选股的好战场",
            [
                "上证50最好的是 valuation_momentum，尤其 Top1/Top3 很亮眼：valuation_momentum Top1 总收益约 118.3%，但这是单票味道，非常浓；Top5 降到约 46.3%，Top10 降到约 37.6%。",
                "这说明上证50的机会更像少数核心资产切换，而不是宽股票池横截面选股。",
                "上证50只有50只股票，且大多流动性充分、机构覆盖充分，大家都看得到。量能和换手差异没有那么强的区分度，所以 liquidity_volume 在上证50里表现弱，这符合交易直觉。",
                "上证50可以做，但交易方式应该不一样：更像风格择时，更像低估值修复 + 趋势确认，更适合和宏观、利率、人民币、北向/机构偏好结合，不适合只靠短期量价横截面硬选。",
                "我的判断：上证50可以作为防守池、核心资产观察池，但不是这套方法的主战场。",
            ],
        ),
        (
            "最重要的横向发现：流动性信号的甜蜜点在中证500",
            [
                "liquidity 这个东西，越往中小盘越重要，但到中证2000又会从优势变成风险。",
                "在上证50里，流动性不是稀缺资源。大家都能买卖，成交差异不够大，所以流动性信号没那么有用。",
                "到中证500，流动性开始变成资金关注度的代理。谁突然有量、谁换手活跃，往往代表资金开始定价新信息，所以它最有效。",
                "到中证2000，流动性仍然有效，但它同时带来另一面：高换手、小市值、情绪票、涨跌停、滑点、隔夜风险。信号更刺激，但交易摩擦也更刺激。",
                "三者关系不是简单线性：上证50流动性差异太小；中证500流动性差异刚刚好；中证2000流动性差异很大，但交易摩擦也很大。",
            ],
        ),
    ]

    for title, paragraphs in sections:
        add_heading(doc, title, level=1)
        for text in paragraphs:
            add_para(doc, text)

    add_heading(doc, "交易优先级", level=1)
    add_table(
        doc,
        ["优先级", "方向", "交易定位"],
        [
            ["1", "中证500 momentum_liquidity Top5/10/15", "最应该继续做成实盘候选系统的方向"],
            ["2", "中证2000 momentum_liquidity Top10/15/20", "战术增强池，先做可交易过滤，再看收益是否还能留下"],
            ["3", "上证50 valuation_momentum Top3/5", "核心资产风格策略，不期待大量横截面机会"],
        ],
    )

    add_heading(doc, "下一步交易真实性检验", level=1)
    for text in [
        "1. 加成本：单边 10bp / 20bp / 30bp / 50bp。",
        "2. 加成交额过滤：例如过去20日均成交额分档。",
        "3. 加 ST 过滤。",
        "4. 加涨跌停可成交过滤。",
        "5. 看每期换手率。",
        "6. 看收益来自哪些月份、哪些票。",
        "7. 看 TopK 扩容后收益是否平滑。",
        "8. 对中证2000单独看极端亏损日和无法退出风险。",
    ]:
        add_para(doc, text)
    add_para(doc, "如果这些做完后，中证500还能保留明显收益，那它就真的有继续推进价值。中证2000如果成本后还能剩一半，也值得作为进攻型副策略。")
    return doc


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    try:
        doc.save(OUT_PATH)
        saved_path = OUT_PATH
    except PermissionError:
        # Word locks open .docx files on Windows, so fall back to a fresh path.
        saved_path = OUT_DIR / "three_index_trader_insights_utf8.docx"
        doc.save(saved_path)
    print(saved_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
