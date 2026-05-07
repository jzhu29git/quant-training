#!/usr/bin/env python3
"""Generate a Chinese Word guide for quant_data/csi500_2y_run."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "quant_data" / "comparison_reports"
OUT_PATH = OUT_DIR / "csi500_2y_run_folder_guide.docx"

FONT_BODY = "Microsoft YaHei"
FONT_TITLE = "Microsoft YaHei"


def set_run_font(run, font_name: str = FONT_BODY, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def set_style_font(style, font_name: str, size: float | None = None) -> None:
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(run, FONT_TITLE, 20, True, (31, 78, 121))
    else:
        set_run_font(run, FONT_TITLE, 14 if level == 1 else 12, True, (31, 78, 121))


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, 10.5)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, 10.5)


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, "Consolas", 9.5)
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.space_after = Pt(4)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        paragraph = table.rows[0].cells[idx].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, FONT_BODY, 9.5, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            run = cells[idx].paragraphs[0].add_run(value)
            set_run_font(run, FONT_BODY, 9)


def build_doc() -> Document:
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title", "Subtitle", "List Bullet"]:
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], FONT_BODY, 10.5)

    add_heading(doc, "中证500研究目录使用指南", level=0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(r"C:\Users\Administrator\quant-trading-cn\quant_data\csi500_2y_run")
    set_run_font(run, FONT_BODY, 10.5, False, (88, 88, 88))

    add_heading(doc, "Executive Summary", level=1)
    add_para(doc, "这个目录是中证500两年研究档案包，不是单一结果文件。它覆盖了从原始数据、特征工程、模型训练、walk-forward回测、因子组诊断、TopK测试，到最新交易信号的完整链路。")
    add_para(doc, "交易执行上，当前建议默认关注 momentum_liquidity 因子组，尤其是 topk_tests/momentum_liquidity/topk_10 和 feature_group_tests/momentum_liquidity/models/inference_scores_latest.parquet。")
    add_para(doc, "如果只是日常使用，最常看的文件是 stock_list.parquet、run_meta.json、topk_tests/topk_results.csv、feature_group_tests/momentum_liquidity/models/inference_scores_latest.parquet，以及 paper trading runbook。")

    add_heading(doc, "目录回答的五个问题", level=1)
    for text in [
        "我研究的是谁：中证500当前500只成分股。",
        "我用了什么数据：过去约两年的日线行情、成交额、换手率、估值数据。",
        "模型学了什么：动量、流动性、估值、波动率等因子对未来5个交易日表现的提示。",
        "历史上效果如何：walk-forward回测、不同因子组、不同TopK组合表现。",
        "现在应该买什么：最新 inference_scores_latest.parquet 给出当前打分和Top picks。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "根目录关键文件", level=1)
    add_table(
        doc,
        ["文件/目录", "作用", "怎么用"],
        [
            ["stock_list.parquet / stock_list.csv", "中证500成分股列表，共500只", "确认股票池是谁"],
            ["run_meta.json", "本次运行元信息", "看日期区间、成分日期、下载成功率"],
            ["download_summary.csv", "每只股票下载情况", "检查是否缺数据"],
            ["ml_features_ready.parquet", "训练/回测用完整特征表，含future_return和label", "模型训练、回测、catch-up历史回放"],
            ["inference_features_latest.parquet", "最新一天推理特征，不含未来收益", "给最新500只股票打分"],
            ["models/", "主模型输出", "看全因子模型和最新Top picks"],
            ["backtest/", "主模型walk-forward回测", "看历史净值、交易日志、OOS预测"],
            ["feature_group_tests/", "不同因子组测试", "判断哪类因子更有效"],
            ["topk_tests/", "Top1/3/5/10/15/20/30测试", "决定持仓数量"],
            ["diagnostics/", "分层收益、Rank IC、相对基准等诊断", "判断信号质量"],
        ],
    )

    add_heading(doc, "本次运行质量", level=1)
    add_code(doc, "universe: csi500")
    add_code(doc, "index source: 000905")
    add_code(doc, "constituent_date: 2026-04-30")
    add_code(doc, "start_date: 2024-04-30")
    add_code(doc, "end_date: 2026-04-30")
    add_code(doc, "stock_count: 500")
    add_code(doc, "download_success_count: 500")
    add_code(doc, "download_failure_count: 0")
    add_para(doc, "交易员视角：这次中证500数据下载是完整的，500只成分股全部成功。")

    add_heading(doc, "原始数据层", level=1)
    add_heading(doc, "daily_kline/", level=2)
    add_para(doc, "每只股票一个parquet文件，例如000009.parquet、600517.parquet、688819.parquet。里面是日线行情，包括open、high、low、close、volume、amount、turnover、pct_chg、change等字段。")
    add_para(doc, "用途：查看单票历史价格、成交量、换手率、涨跌幅；也是所有动量和流动性因子的来源。")
    add_heading(doc, "daily_valuation/", level=2)
    add_para(doc, "每只股票一个parquet文件，保存估值相关字段，例如pe_ttm、pb、ps、pcf等。")
    add_para(doc, "用途：用于估值因子和 valuation_momentum 组合。")

    add_heading(doc, "特征层", level=1)
    add_heading(doc, "ml_features_ready.parquet", level=2)
    add_para(doc, "训练和回测用的大表。本次约228,919行，覆盖500只股票、460个交易日。它包含历史特征、未来5日收益future_return、以及label。")
    add_para(doc, "future_return 是未来5个交易日收益，label 是未来5日是否超过阈值，例如2%。这两个字段只能用于训练和回测，不能用于真实当天选股。")
    add_heading(doc, "inference_features_latest.parquet", level=2)
    add_para(doc, "最新一天的推理特征。本次是2026-04-30，500行。它没有future_return和label，因为真实交易时未来还没发生。")
    add_para(doc, "用途：每次最新打分和paper trading调仓之前，都会先生成或更新这个文件。")

    add_heading(doc, "模型层", level=1)
    add_para(doc, "models/目录是主模型输出，包含 lightgbm_model.txt、training_metadata.json、feature_importance.csv、inference_scores_latest.parquet。")
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["lightgbm_model.txt", "训练好的LightGBM模型"],
            ["training_metadata.json", "训练参数、训练/验证区间、指标"],
            ["feature_importance.csv", "特征重要性"],
            ["inference_scores_latest.parquet", "最新股票打分结果"],
        ],
    )
    add_para(doc, "主模型最新Top10来自全因子模型，但交易执行上我们当前更偏向 momentum_liquidity 因子组，因为它在中证500上的交易形态更干净。")

    add_heading(doc, "主回测层", level=1)
    add_para(doc, "backtest/目录保存主模型walk-forward结果，包括summary.json、equity_curve.parquet、trade_log.parquet、oos_predictions.parquet。")
    add_code(doc, "backtest_start: 2025-06-17")
    add_code(doc, "backtest_end: 2026-04-21")
    add_code(doc, "num_rebalances: 42")
    add_code(doc, "top_k: 5")
    add_code(doc, "rebalance_every: 5")
    add_code(doc, "retrain_every: 20")
    add_code(doc, "portfolio_total_return: 72.56%")
    add_code(doc, "max_drawdown: -11.07%")
    add_code(doc, "win_rate: 59.52%")
    add_para(doc, "交易员视角：历史OOS阶段每5天换一次、持有Top5，收益表现不错，最大回撤约11%。但这是历史回测，不是实盘承诺。")

    add_heading(doc, "因子组测试层", level=1)
    add_para(doc, "feature_group_tests/目录用于回答：到底是哪类因子有效？")
    add_table(
        doc,
        ["因子组", "含义"],
        [
            ["all_features", "全部可用特征"],
            ["momentum_liquidity", "动量 + 流动性，当前交易上最关注"],
            ["valuation_momentum", "估值 + 动量"],
            ["momentum_trend", "趋势和短中期动量"],
            ["liquidity_volume", "成交量、成交额、换手"],
            ["valuation", "估值因子"],
            ["volatility_position", "波动率和价格位置"],
            ["price_ohlc", "OHLC价格类特征"],
        ],
    )
    add_para(doc, "momentum_liquidity主要使用ma5、ma20、bias_20、pct_chg_5d、pct_chg_20d、pct_chg、change、volume、amount、turnover、turnover_ma5、volume_ma5。")

    add_heading(doc, "当前更推荐看的最新Top10", level=1)
    add_para(doc, "路径：feature_group_tests/momentum_liquidity/models/inference_scores_latest.parquet")
    add_table(
        doc,
        ["排名", "代码", "名称", "说明"],
        [
            ["1", "688819", "天能股份", "momentum_liquidity 当前最高分"],
            ["2", "601567", "三星电气", ""],
            ["3", "688582", "芯动联科", ""],
            ["4", "002432", "九安医疗", ""],
            ["5", "000537", "绿发电力", ""],
            ["6", "600032", "浙江新能", ""],
            ["7", "002624", "完美世界", ""],
            ["8", "688709", "成都华微", ""],
            ["9", "300058", "蓝色光标", ""],
            ["10", "600562", "国睿科技", ""],
        ],
    )

    add_heading(doc, "TopK测试层", level=1)
    add_para(doc, "topk_tests/目录用于回答：买Top1、Top3、Top5、Top10还是Top20？每个目录都有summary.json、equity_curve.parquet、trade_log.parquet、oos_predictions.parquet。")
    add_para(doc, "从现有结果看，中证500的 momentum_liquidity Top5 历史表现非常亮眼，总收益约147.35%，最大回撤约-10.99%。但Top5更集中，真实执行波动也更大。因此paper trading默认使用Top10，更稳一点。")

    add_heading(doc, "diagnostics诊断层", level=1)
    add_para(doc, "diagnostics/目录用于判断信号质量，不直接给买卖列表。重点看quantile曲线、Rank IC、benchmark excess等。")
    add_bullet(doc, "quantile_5_curve.csv / quantile_10_curve.csv：看高分组是否持续跑赢低分组。")
    add_bullet(doc, "rank_ic_timeseries.csv：看模型分数和未来收益的排序相关性。")
    add_bullet(doc, "benchmark_excess.csv：看相对基准是否有超额。")
    add_bullet(doc, "diagnostics_summary.json：诊断摘要。")

    add_heading(doc, "日常使用路线", level=1)
    add_table(
        doc,
        ["目标", "看什么文件", "动作"],
        [
            ["确认数据完整性", "run_meta.json / download_summary.csv", "看是否500只全部成功"],
            ["看最新候选股", "feature_group_tests/momentum_liquidity/models/inference_scores_latest.parquet", "按score排序看Top10"],
            ["看历史策略是否靠谱", "topk_tests/topk_results.csv", "比较group和top_k"],
            ["看历史每次买了什么", "topk_tests/momentum_liquidity/topk_10/trade_log.parquet", "检查选股路径"],
            ["看净值曲线", "topk_tests/momentum_liquidity/topk_10/equity_curve.parquet", "看收益和回撤"],
            ["做paper trading", "paper_trading_config.yaml / RUNBOOK", "按默认流程dry-run再成交"],
        ],
    )

    add_heading(doc, "常用命令", level=1)
    add_para(doc, "先看计划，不成交：")
    add_code(doc, r"python scripts\run_incremental_rebalance.py --index csi500 --update-data --dry-run --force")
    add_para(doc, "确认后正式paper rebalance：")
    add_code(doc, r"python scripts\run_incremental_rebalance.py --index csi500 --update-data --force")
    add_para(doc, "如果隔了很多天，先做catch-up研究回放，不改broker：")
    add_code(doc, r"python scripts\run_catchup_rebalance.py --index csi500 --from-date YYYY-MM-DD --to-date YYYY-MM-DD")

    add_heading(doc, "一句话总结", level=1)
    add_para(doc, "csi500_2y_run 是中证500策略的完整研究档案。研究时看 topk_tests 和 diagnostics，交易时看 momentum_liquidity 最新 scores 和 paper trading runbook。")

    return doc


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_PATH)
    print(OUT_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
