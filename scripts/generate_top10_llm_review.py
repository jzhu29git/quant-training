#!/usr/bin/env python3
"""Generate an LLM-enhanced TopK trader review.

The script never stores API keys. Providers are configured in
llm_review_config.yaml and keys are read from environment variables.
"""

from __future__ import annotations

import argparse
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
import json
import os
import re
import subprocess
import sys
from pathlib import Path
from typing import Any

import pandas as pd
import requests
import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_PAPER_CONFIG = ROOT / "paper_trading_config.yaml"
DEFAULT_LLM_CONFIG = ROOT / "llm_review_config.yaml"
DEFAULT_ENV_PATH = ROOT / ".env"
OUT_DIR = ROOT / "quant_data" / "comparison_reports"
FONT_BODY = "Microsoft YaHei"
FONT_TITLE = "Microsoft YaHei"


def read_yaml(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as handle:
        data = yaml.safe_load(handle) or {}
    if not isinstance(data, dict):
        raise ValueError(f"config must be a mapping: {path}")
    return data


def load_dotenv(path: Path = DEFAULT_ENV_PATH) -> None:
    if not path.exists():
        return
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def get_api_key(env_name: str) -> str:
    value = os.environ.get(env_name, "").strip()
    if value:
        return value
    if os.name != "nt":
        return ""
    try:
        import subprocess

        script = (
            f"[Environment]::GetEnvironmentVariable('{env_name}', 'User'); "
            f"[Environment]::GetEnvironmentVariable('{env_name}', 'Machine')"
        )
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", script],
            capture_output=True,
            text=True,
            timeout=10,
        )
        for line in result.stdout.splitlines():
            candidate = line.strip()
            if candidate:
                os.environ[env_name] = candidate
                return candidate
    except Exception:
        return ""
    return ""


def deep_get(payload: dict[str, Any], keys: list[str], default: Any = None) -> Any:
    cur: Any = payload
    for key in keys:
        if not isinstance(cur, dict) or key not in cur:
            return default
        cur = cur[key]
    return cur


def ensure_non_llm_review(index_name: str, group: str, top_k: int, paper_config: Path, output_dir: Path) -> Path:
    existing = output_dir / f"{index_name}_{group}_top{top_k}_non_llm_review.csv"
    if existing.exists():
        return existing
    cmd = [
        sys.executable,
        "scripts/generate_top10_non_llm_review.py",
        "--config",
        str(paper_config),
        "--index",
        index_name,
        "--group",
        group,
        "--top-k",
        str(top_k),
        "--output-dir",
        str(output_dir),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)
    return output_dir / f"{index_name}_{group}_top{top_k}_non_llm_review.csv"


def extract_json(text: str) -> dict[str, Any]:
    text = str(text or "").strip()
    if not text:
        return {}
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    match = re.search(r"\{.*\}", text, flags=re.S)
    if not match:
        return {"raw_text": text}
    try:
        return json.loads(match.group(0))
    except json.JSONDecodeError:
        return {"raw_text": text}


def review_prompt(row: dict[str, Any]) -> tuple[str, str]:
    system = (
        "你是A股交易员的风控复核助手。你只基于用户提供的本地量化指标做判断，"
        "不要编造新闻、公告或基本面事实。输出必须是JSON，不要Markdown。"
    )
    user_payload = {
        "task": "复核模型TopK候选股，给出交易员视角保留/谨慎/剔除建议。",
        "decision_values": ["保留", "谨慎", "剔除"],
        "required_json_schema": {
            "llm_decision": "保留|谨慎|剔除",
            "confidence": "0到1之间的小数",
            "trader_summary": "一句话交易员点评",
            "key_risks": ["风险1", "风险2"],
            "watch_items": ["后续观察点1", "后续观察点2"],
            "suggested_action": "具体动作，例如保留但降权/等待承接/剔除",
        },
        "candidate": row,
        "rules": [
            "不要因为模型分数高就无条件保留。",
            "成交额、换手、短线涨跌幅、是否过热/大跌，都要纳入判断。",
            "如果信息不足，请明确说信息不足，不要编造外部事实。",
            "这是paper trading复核，不是投资建议。",
        ],
    }
    return system, json.dumps(user_payload, ensure_ascii=False, indent=2)


def call_chat_completions(provider: dict[str, Any], api_key: str, system: str, user: str, timeout: int) -> dict[str, Any]:
    base_url = str(provider["base_url"]).rstrip("/")
    url = f"{base_url}/chat/completions"
    payload = {
        "model": provider["model"],
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": float(provider.get("temperature", 0.2)),
        "response_format": {"type": "json_object"},
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    response = requests.post(url, headers=headers, json=payload, timeout=timeout)
    response.raise_for_status()
    data = response.json()
    content = data["choices"][0]["message"]["content"]
    return extract_json(content)


def call_responses(provider: dict[str, Any], api_key: str, system: str, user: str, timeout: int, max_output_tokens: int) -> dict[str, Any]:
    base_url = str(provider["base_url"]).rstrip("/")
    url = f"{base_url}/responses"
    payload = {
        "model": provider["model"],
        "instructions": system,
        "input": user,
        "temperature": float(provider.get("temperature", 0.2)),
        "max_output_tokens": int(max_output_tokens),
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    response = requests.post(url, headers=headers, json=payload, timeout=timeout)
    response.raise_for_status()
    data = response.json()
    if "output_text" in data:
        return extract_json(data["output_text"])
    chunks: list[str] = []
    for item in data.get("output", []):
        for content in item.get("content", []):
            if "text" in content:
                chunks.append(str(content["text"]))
    return extract_json("\n".join(chunks))


def call_provider(provider: dict[str, Any], row: dict[str, Any], review_cfg: dict[str, Any]) -> dict[str, Any]:
    api_key_env = str(provider.get("api_key_env") or "")
    api_key = get_api_key(api_key_env)
    if not api_key:
        raise RuntimeError(f"missing API key env var: {api_key_env}")
    system, user = review_prompt(row)
    timeout = int(review_cfg.get("timeout_seconds") or 60)
    max_output_tokens = int(review_cfg.get("max_output_tokens") or 900)
    wire_api = str(provider.get("wire_api") or "chat_completions")
    if wire_api == "responses":
        return call_responses(provider, api_key, system, user, timeout, max_output_tokens)
    return call_chat_completions(provider, api_key, system, user, timeout)


def provider_keys_from_args(value: str, llm_cfg: dict[str, Any]) -> list[str]:
    if value.strip():
        return [item.strip() for item in value.split(",") if item.strip()]
    default = str(llm_cfg.get("default_provider") or "deepseek")
    return [default]


def normalize_decision(value: Any) -> str:
    text = str(value or "").strip()
    if "剔除" in text:
        return "剔除"
    if "谨慎" in text or "觀望" in text or "观望" in text:
        return "谨慎"
    if "保留" in text:
        return "保留"
    return "未返回"


def call_one_provider(
    *,
    provider_key: str,
    provider: dict[str, Any],
    payload: dict[str, Any],
    review_cfg: dict[str, Any],
) -> dict[str, Any]:
    try:
        result = call_provider(provider, payload, review_cfg)
        decision = normalize_decision(result.get("llm_decision") or result.get("decision"))
        return {
            "rank": payload.get("rank"),
            "code": payload.get("code"),
            "name": payload.get("name"),
            "provider": provider_key,
            "provider_name": provider.get("name") or provider_key,
            "status": "ok",
            "llm_decision": decision,
            "confidence": result.get("confidence", ""),
            "trader_summary": result.get("trader_summary", result.get("summary", "")),
            "key_risks": "；".join(result.get("key_risks", [])) if isinstance(result.get("key_risks"), list) else str(result.get("key_risks", "")),
            "watch_items": "；".join(result.get("watch_items", [])) if isinstance(result.get("watch_items"), list) else str(result.get("watch_items", "")),
            "suggested_action": result.get("suggested_action", ""),
            "llm_raw_json": json.dumps(result, ensure_ascii=False),
        }
    except Exception as exc:
        return {
            "rank": payload.get("rank"),
            "code": payload.get("code"),
            "name": payload.get("name"),
            "provider": provider_key,
            "provider_name": provider.get("name") or provider_key,
            "status": "error",
            "llm_decision": "未返回",
            "confidence": "",
            "trader_summary": "",
            "key_risks": "",
            "watch_items": "",
            "suggested_action": "",
            "llm_raw_json": json.dumps({"error": str(exc)}, ensure_ascii=False),
        }


def run_multi_provider_review(
    frame: pd.DataFrame,
    *,
    provider_keys: list[str],
    llm_cfg: dict[str, Any],
    review_cfg: dict[str, Any],
    max_workers: int,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    providers = deep_get(llm_cfg, ["providers"], {}) or {}
    selected: dict[str, dict[str, Any]] = {}
    for key in provider_keys:
        provider = providers.get(key)
        if not provider:
            raise RuntimeError(f"unknown provider: {key}")
        selected[key] = provider

    tasks = []
    with ThreadPoolExecutor(max_workers=max(1, int(max_workers))) as executor:
        for _, row in frame.iterrows():
            payload = row.where(pd.notna(row), None).to_dict()
            for provider_key, provider in selected.items():
                tasks.append(
                    executor.submit(
                        call_one_provider,
                        provider_key=provider_key,
                        provider=provider,
                        payload=payload,
                        review_cfg=review_cfg,
                    )
                )
        provider_rows = [future.result() for future in as_completed(tasks)]

    long_df = pd.DataFrame(provider_rows).sort_values(["rank", "provider"]).reset_index(drop=True)
    summary_rows: list[dict[str, Any]] = []
    base_lookup = {str(row.code): row._asdict() for row in frame.itertuples(index=False)}
    for code, group in long_df.groupby("code", sort=False):
        base = base_lookup.get(str(code), {})
        ok = group[group["status"] == "ok"].copy()
        votes = Counter(ok["llm_decision"].map(normalize_decision).tolist())
        if votes:
            final_decision, final_votes = votes.most_common(1)[0]
        else:
            final_decision, final_votes = "未返回", 0
        valid_vote_count = int(sum(votes.values()))
        disagreement = len([key for key, count in votes.items() if key != "未返回" and count > 0]) > 1
        summaries = [f"{row.provider}:{row.llm_decision}-{row.trader_summary}" for row in ok.itertuples(index=False)]
        risks = [str(row.key_risks) for row in ok.itertuples(index=False) if str(row.key_risks or "").strip()]
        actions = [str(row.suggested_action) for row in ok.itertuples(index=False) if str(row.suggested_action or "").strip()]
        summary_rows.append(
            {
                **base,
                "provider_count": len(selected),
                "successful_provider_count": int(len(ok)),
                "vote保留": int(votes.get("保留", 0)),
                "vote谨慎": int(votes.get("谨慎", 0)),
                "vote剔除": int(votes.get("剔除", 0)),
                "final_decision": final_decision,
                "final_vote_share": round(final_votes / valid_vote_count, 4) if valid_vote_count else 0.0,
                "disagreement": bool(disagreement),
                "provider_summaries": " | ".join(summaries),
                "merged_risks": "；".join(risks),
                "merged_actions": "；".join(actions),
            }
        )
    summary_df = pd.DataFrame(summary_rows).sort_values("rank").reset_index(drop=True)
    return long_df, summary_df


def set_run_font(run, font_name: str = FONT_BODY, size: float | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def set_style_font(style, font_name: str, size: float | None = None) -> None:
    style.font.name = font_name
    if size is not None:
        style.font.size = Pt(size)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    if level == 0:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(run, FONT_TITLE, 18, True, (31, 78, 121))
    else:
        set_run_font(run, FONT_TITLE, 13 if level == 1 else 11.5, True, (31, 78, 121))


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, 10.5)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, FONT_BODY, 8.8, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            run = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(run, FONT_BODY, 8.2)


def build_doc(
    frame: pd.DataFrame,
    long_frame: pd.DataFrame,
    *,
    index_name: str,
    group: str,
    provider_name: str,
    csv_path: Path,
    long_csv_path: Path,
) -> Document:
    doc = Document()
    for style_name in ["Normal", "Heading 1", "Heading 2", "Title"]:
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], FONT_BODY, 10.5)
    add_heading(doc, f"{index_name.upper()} Top10 LLM交易员复核", level=0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signal_date = str(frame["date"].iloc[0])[:10] if "date" in frame.columns and not frame.empty else ""
    run = subtitle.add_run(f"因子组：{group} | Providers：{provider_name} | 信号日期：{signal_date}")
    set_run_font(run, FONT_BODY, 10.5, False, (88, 88, 88))

    add_heading(doc, "Executive Summary", level=1)
    counts = frame["final_decision"].fillna("未返回").value_counts().to_dict()
    disagreement_count = int(frame["disagreement"].fillna(False).sum()) if "disagreement" in frame.columns else 0
    add_para(doc, f"本报告在无LLM规则复核基础上，加入多模型并行复核和投票。本轮最终建议：{counts}；出现模型分歧 {disagreement_count} 只。")
    add_para(doc, "LLM只基于本地指标做复核，不包含实时新闻/公告，不能替代人工风控和paper trading dry-run。最终执行前仍建议先dry-run。")

    add_heading(doc, "投票汇总表", level=1)
    rows = []
    for _, row in frame.iterrows():
        rows.append(
            [
                row.get("rank", ""),
                row.get("code", ""),
                row.get("name", ""),
                row.get("decision", ""),
                f"保留{row.get('vote保留', 0)}/谨慎{row.get('vote谨慎', 0)}/剔除{row.get('vote剔除', 0)}",
                row.get("final_decision", ""),
                f"{float(row.get('final_vote_share', 0) or 0) * 100:.0f}%",
                "是" if bool(row.get("disagreement", False)) else "否",
                row.get("merged_actions", ""),
            ]
        )
    add_table(doc, ["Rank", "代码", "名称", "规则建议", "投票", "最终建议", "一致度", "分歧", "综合动作"], rows)

    if disagreement_count:
        add_heading(doc, "分歧提醒", level=1)
        for _, row in frame[frame["disagreement"] == True].iterrows():  # noqa: E712
            add_para(doc, f"{row.get('rank')}. {row.get('code')} {row.get('name')}：模型投票存在分歧，建议人工复核后再进入paper组合。")

    add_heading(doc, "逐只详情", level=1)
    for _, row in frame.iterrows():
        add_heading(doc, f"{row.get('rank')}. {row.get('code')} {row.get('name')} | 最终：{row.get('final_decision')}", level=2)
        add_para(doc, f"规则风险：{row.get('risk_flags', '')}")
        add_para(doc, f"多模型风险合并：{row.get('merged_risks', '')}")
        add_para(doc, f"多模型动作合并：{row.get('merged_actions', '')}")
        provider_details = long_frame[long_frame["code"].astype(str) == str(row.get("code"))]
        for detail in provider_details.itertuples(index=False):
            add_para(doc, f"{detail.provider}：{detail.llm_decision}；{detail.trader_summary}；动作：{detail.suggested_action}")

    add_heading(doc, "输出文件", level=1)
    add_para(doc, f"投票汇总CSV：{csv_path}")
    add_para(doc, f"逐模型明细CSV：{long_csv_path}")
    return doc


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate LLM-enhanced TopK review.")
    parser.add_argument("--paper-config", default=str(DEFAULT_PAPER_CONFIG))
    parser.add_argument("--llm-config", default=str(DEFAULT_LLM_CONFIG))
    parser.add_argument("--index", choices=["csi500", "csi2000", "sse50"], default="csi500")
    parser.add_argument("--group", default="")
    parser.add_argument("--top-k", type=int, default=0)
    parser.add_argument("--provider", default="")
    parser.add_argument("--providers", default="", help="Comma-separated provider keys. Overrides --provider.")
    parser.add_argument("--output-dir", default=str(OUT_DIR))
    parser.add_argument("--limit", type=int, default=0, help="Only review first N rows; useful for provider testing.")
    parser.add_argument("--max-workers", type=int, default=4)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    load_dotenv()
    paper_cfg = read_yaml(Path(args.paper_config))
    llm_cfg = read_yaml(Path(args.llm_config))
    index_cfg = deep_get(paper_cfg, ["indices", args.index], {}) or {}
    group = args.group or str(index_cfg.get("group") or "momentum_liquidity")
    top_k = args.top_k or int(index_cfg.get("top_k") or 10)
    provider_arg = args.providers or args.provider
    provider_keys = provider_keys_from_args(provider_arg, llm_cfg)
    providers_cfg = deep_get(llm_cfg, ["providers"], {}) or {}
    for provider_key in provider_keys:
        if provider_key not in providers_cfg:
            raise RuntimeError(f"unknown provider: {provider_key}")
    review_cfg = deep_get(llm_cfg, ["review"], {}) or {}

    out_dir = Path(args.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    non_llm_csv = ensure_non_llm_review(args.index, group, top_k, Path(args.paper_config), out_dir)
    frame = pd.read_csv(non_llm_csv, encoding="utf-8-sig")
    if args.limit and args.limit > 0:
        frame = frame.head(args.limit).copy()

    long_frame, summary_frame = run_multi_provider_review(
        frame,
        provider_keys=provider_keys,
        llm_cfg=llm_cfg,
        review_cfg=review_cfg,
        max_workers=args.max_workers,
    )
    provider_slug = "_".join(provider_keys)
    suffix = f"{args.index}_{group}_top{top_k}_{provider_slug}_llm_vote_review"
    csv_path = out_dir / f"{suffix}.csv"
    long_csv_path = out_dir / f"{suffix}_provider_details.csv"
    docx_path = out_dir / f"{suffix}.docx"
    summary_frame.to_csv(csv_path, index=False, encoding="utf-8-sig")
    long_frame.to_csv(long_csv_path, index=False, encoding="utf-8-sig")
    provider_name = ", ".join(str(providers_cfg[key].get("name") or key) for key in provider_keys)
    doc = build_doc(
        summary_frame,
        long_frame,
        index_name=args.index,
        group=group,
        provider_name=provider_name,
        csv_path=csv_path,
        long_csv_path=long_csv_path,
    )
    doc.save(docx_path)
    print(
        json.dumps(
            {
                "csv": str(csv_path),
                "provider_details_csv": str(long_csv_path),
                "docx": str(docx_path),
                "rows": len(summary_frame),
                "providers": provider_keys,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
